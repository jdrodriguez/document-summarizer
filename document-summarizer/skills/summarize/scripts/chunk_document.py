#!/usr/bin/env python3
"""
Document chunking script for the document-summarizer skill.

Extracts text from PDF, DOCX, TXT, and MD files, detects section boundaries,
and splits into manageable chunks for parallel LLM summarization.
Supports single files or directories containing multiple documents.

Usage:
    python3 chunk_document.py <input_path> <output_dir> [--max-tokens 4000] [--overlap 200]

    input_path: a single file (.pdf, .docx, .txt, .md) OR a directory of files
"""
import argparse
import json
import os
import re
import subprocess
import shutil
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Dependency imports (with graceful fallback messages)
# ---------------------------------------------------------------------------
try:
    import tiktoken
    ENCODER = tiktoken.get_encoding("cl100k_base")
    def count_tokens(text: str) -> int:
        return len(ENCODER.encode(text))
except Exception:
    # tiktoken may fail on import OR when downloading its encoding file
    # (e.g. in sandboxed/offline environments like Cowork's VM)
    ENCODER = None
    def count_tokens(text: str) -> int:
        return len(text) // 4  # ~1 token per 4 chars

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}

# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------
class TextBlock:
    """A block of extracted text with metadata."""
    __slots__ = ("text", "page", "heading_level", "style", "confidence")

    def __init__(self, text: str, page: int = 0, heading_level: int = 0,
                 style: str = "", confidence: float = 0.0):
        self.text = text
        self.page = page
        self.heading_level = heading_level
        self.style = style
        self.confidence = confidence

class SectionBoundary:
    """A detected section boundary."""
    __slots__ = ("index", "heading", "level", "confidence", "page")

    def __init__(self, index: int, heading: str, level: int,
                 confidence: float, page: int):
        self.index = index
        self.heading = heading
        self.level = level
        self.confidence = confidence
        self.page = page

# ---------------------------------------------------------------------------
# Text extraction — PDF
# ---------------------------------------------------------------------------
def extract_pdf_pymupdf(filepath: str) -> list[TextBlock]:
    if fitz is None:
        return []
    blocks = []
    try:
        doc = fitz.open(filepath)
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                blocks.append(TextBlock(text=text, page=page_num + 1))
        doc.close()
    except Exception as e:
        print(f"PyMuPDF extraction failed: {e}", file=sys.stderr)
        return []
    return blocks


def extract_pdf_pdfplumber(filepath: str) -> list[TextBlock]:
    if pdfplumber is None:
        return []
    blocks = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    blocks.append(TextBlock(text=text, page=i + 1))
    except Exception as e:
        print(f"pdfplumber extraction failed: {e}", file=sys.stderr)
        return []
    return blocks


def extract_pdf_pdftotext(filepath: str) -> list[TextBlock]:
    if not shutil.which("pdftotext"):
        return []
    blocks = []
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", filepath, "-"],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and result.stdout.strip():
            pages = result.stdout.split("\f")
            for i, page_text in enumerate(pages):
                if page_text.strip():
                    blocks.append(TextBlock(text=page_text, page=i + 1))
    except Exception as e:
        print(f"pdftotext extraction failed: {e}", file=sys.stderr)
        return []
    return blocks


def extract_pdf(filepath: str) -> list[TextBlock]:
    """Extract text from PDF with cascading fallbacks."""
    for extractor in (extract_pdf_pymupdf, extract_pdf_pdfplumber, extract_pdf_pdftotext):
        blocks = extractor(filepath)
        if sum(len(b.text) for b in blocks) > 100:
            return blocks
    return blocks

# ---------------------------------------------------------------------------
# Text extraction — DOCX
# ---------------------------------------------------------------------------
def extract_docx(filepath: str) -> list[TextBlock]:
    if DocxDocument is None:
        print("python-docx not available", file=sys.stderr)
        return []

    blocks = []
    try:
        doc = DocxDocument(filepath)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                blocks.append(TextBlock(text="\n", style="blank"))
                continue

            style_name = para.style.name if para.style else ""
            heading_level = 0
            confidence = 0.0

            if style_name.startswith("Heading"):
                try:
                    heading_level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    heading_level = 1
                confidence = 1.0
            elif style_name in ("Title", "Subtitle"):
                heading_level = 1 if style_name == "Title" else 2
                confidence = 1.0

            blocks.append(TextBlock(
                text=text, heading_level=heading_level,
                style=style_name, confidence=confidence,
            ))

        for table in doc.tables:
            table_text = "[TABLE]\n"
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_text += " | ".join(cells) + "\n"
            table_text += "[/TABLE]\n"
            blocks.append(TextBlock(text=table_text, style="table"))

    except Exception as e:
        print(f"DOCX extraction failed: {e}", file=sys.stderr)
        return []
    return blocks

# ---------------------------------------------------------------------------
# Text extraction — Plain text and Markdown
# ---------------------------------------------------------------------------
MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)")

def extract_text_file(filepath: str) -> list[TextBlock]:
    """Extract text from .txt or .md files, detecting markdown headings."""
    blocks = []
    ext = Path(filepath).suffix.lower()
    is_markdown = ext in (".md", ".markdown")

    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            for line in f:
                text = line.rstrip("\n")

                if not text.strip():
                    blocks.append(TextBlock(text="\n", style="blank"))
                    continue

                heading_level = 0
                confidence = 0.0

                # Detect markdown headings
                if is_markdown:
                    m = MD_HEADING_RE.match(text)
                    if m:
                        heading_level = len(m.group(1))
                        confidence = 1.0

                blocks.append(TextBlock(
                    text=text, heading_level=heading_level,
                    style="markdown" if is_markdown else "plain",
                    confidence=confidence,
                ))
    except Exception as e:
        print(f"Text file extraction failed: {e}", file=sys.stderr)
        return []
    return blocks

# ---------------------------------------------------------------------------
# Unified extraction dispatcher
# ---------------------------------------------------------------------------
def extract_file(filepath: str) -> list[TextBlock]:
    """Route extraction based on file extension."""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return extract_pdf(filepath)
    elif ext == ".docx":
        return extract_docx(filepath)
    elif ext in (".txt", ".md", ".markdown"):
        return extract_text_file(filepath)
    else:
        print(f"Unsupported file type: {ext}", file=sys.stderr)
        return []

# ---------------------------------------------------------------------------
# Section boundary detection
# ---------------------------------------------------------------------------
SECTION_PATTERNS = [
    (re.compile(r"^\s*(\d+\.)+\s+\S"), 0.9),
    (re.compile(r"^\s*[IVXLCDM]+\.\s+\S"), 0.85),
    (re.compile(r"^\s*[A-Z]\.\s+\S"), 0.8),
    (re.compile(r"^\s*(Section|Article|Chapter|Part)\s+[\dIVXLCDM]+", re.IGNORECASE), 0.95),
    (re.compile(r"^\s*(ARTICLE|SECTION|CHAPTER|PART)\s+[\dIVXLCDM]+"), 0.95),
]


def clean_heading(text: str) -> str:
    """Normalize heading text: collapse whitespace, strip page numbers."""
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'^[A-Z]?-?\d+\s+', '', text).strip()
    # Strip leading markdown hashes
    text = re.sub(r'^#+\s*', '', text).strip()
    return text or "Untitled Section"


def detect_boundaries(blocks: list[TextBlock]) -> list[SectionBoundary]:
    """Detect section boundaries in a list of text blocks."""
    boundaries = []
    blank_count = 0

    for i, block in enumerate(blocks):
        text = block.text.strip()

        # Already marked as heading (DOCX style or markdown heading)
        if block.confidence >= 0.9 and block.heading_level > 0:
            boundaries.append(SectionBoundary(
                index=i, heading=clean_heading(text), level=block.heading_level,
                confidence=block.confidence, page=block.page,
            ))
            blank_count = 0
            continue

        if not text:
            blank_count += 1
            continue

        best_confidence = 0.0
        detected_level = 1

        for pattern, conf in SECTION_PATTERNS:
            if pattern.match(text):
                best_confidence = max(best_confidence, conf)
                dot_match = re.match(r"^\s*([\d.]+)", text)
                if dot_match:
                    depth = dot_match.group(1).rstrip(".").count(".") + 1
                    detected_level = min(depth, 4)
                break

        # ALL-CAPS short lines (likely headers)
        if (not best_confidence and len(text) < 120
                and len(text.split()) >= 2
                and text == text.upper()
                and not text.startswith("[TABLE")):
            best_confidence = 0.7
            detected_level = 1

        if blank_count >= 3 and best_confidence == 0:
            best_confidence = 0.3
            detected_level = 1

        if best_confidence > 0:
            boundaries.append(SectionBoundary(
                index=i, heading=clean_heading(text[:120]),
                level=detected_level, confidence=best_confidence,
                page=block.page,
            ))

        blank_count = 0

    return boundaries

# ---------------------------------------------------------------------------
# Chunking strategies
# ---------------------------------------------------------------------------
def blocks_to_text(blocks: list[TextBlock]) -> str:
    return "\n".join(b.text for b in blocks)


def split_text_by_tokens(text: str, max_tokens: int, overlap: int) -> list[str]:
    """Split text into overlapping chunks at sentence boundaries."""
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks = []
    current = []
    current_tokens = 0

    for sentence in sentences:
        stokens = count_tokens(sentence)
        if current_tokens + stokens > max_tokens and current:
            chunks.append(" ".join(current))
            overlap_parts = []
            overlap_tokens = 0
            for s in reversed(current):
                st = count_tokens(s)
                if overlap_tokens + st > overlap:
                    break
                overlap_parts.insert(0, s)
                overlap_tokens += st
            current = overlap_parts
            current_tokens = overlap_tokens
        current.append(sentence)
        current_tokens += stokens

    if current:
        chunks.append(" ".join(current))

    return chunks if chunks else [text]


def chunk_by_structure(blocks, boundaries, max_tokens, overlap):
    """Split text into chunks along section boundaries."""
    if not boundaries:
        return []

    boundaries = sorted(boundaries, key=lambda b: b.index)

    sections = []
    for i, boundary in enumerate(boundaries):
        start_idx = boundary.index
        end_idx = boundaries[i + 1].index if i + 1 < len(boundaries) else len(blocks)
        section_blocks = blocks[start_idx:end_idx]
        section_text = blocks_to_text(section_blocks)
        pages = [b.page for b in section_blocks if b.page > 0]

        sections.append({
            "heading": boundary.heading,
            "level": boundary.level,
            "text": section_text,
            "tokens": count_tokens(section_text),
            "start_page": min(pages) if pages else 0,
            "end_page": max(pages) if pages else 0,
        })

    # Capture preamble text before first boundary
    if boundaries[0].index > 0:
        preamble_blocks = blocks[:boundaries[0].index]
        preamble_text = blocks_to_text(preamble_blocks)
        if preamble_text.strip():
            pages = [b.page for b in preamble_blocks if b.page > 0]
            sections.insert(0, {
                "heading": "Preamble", "level": 0,
                "text": preamble_text,
                "tokens": count_tokens(preamble_text),
                "start_page": min(pages) if pages else 0,
                "end_page": max(pages) if pages else 0,
            })

    # Merge small sections and split large ones
    merged = []
    buffer = None

    for section in sections:
        if section["tokens"] > max_tokens:
            prepend_text = ""
            if buffer and buffer["tokens"] < 100:
                prepend_text = buffer["text"] + "\n\n"
                buffer = None
            elif buffer:
                merged.append(buffer)
                buffer = None
            full_text = prepend_text + section["text"]
            sub_chunks = split_text_by_tokens(full_text, max_tokens, overlap)
            for j, sub_text in enumerate(sub_chunks):
                merged.append({
                    "heading": section["heading"] + (f" (part {j+1})" if len(sub_chunks) > 1 else ""),
                    "level": section["level"],
                    "text": sub_text,
                    "tokens": count_tokens(sub_text),
                    "start_page": section["start_page"],
                    "end_page": section["end_page"],
                })
        elif buffer and (buffer["tokens"] < 100 or buffer["tokens"] + section["tokens"] < max_tokens * 0.8):
            buffer["text"] += "\n\n" + section["text"]
            buffer["tokens"] += section["tokens"]
            if buffer["heading"] != "Preamble" or section["heading"] != "Preamble":
                buffer["heading"] = section["heading"] if buffer["tokens"] < 100 else buffer["heading"] + " / " + section["heading"]
            buffer["end_page"] = max(buffer["end_page"], section["end_page"])
        else:
            if buffer:
                merged.append(buffer)
            buffer = dict(section)

    if buffer:
        merged.append(buffer)

    chunks = []
    for i, section in enumerate(merged):
        chunks.append({
            "id": i + 1,
            "heading": section["heading"],
            "level": section["level"],
            "text": section["text"],
            "tokens": section["tokens"],
            "start_page": section["start_page"],
            "end_page": section["end_page"],
        })
    return chunks


def chunk_by_tokens(blocks, max_tokens, overlap):
    """Fallback: split all text by token count with overlap."""
    full_text = blocks_to_text(blocks)
    sub_texts = split_text_by_tokens(full_text, max_tokens, overlap)
    total_pages = max((b.page for b in blocks if b.page > 0), default=1)
    total_chars = max(len(full_text), 1)
    pages_per_chunk = max(1, total_pages // max(len(sub_texts), 1))

    chunks = []
    char_offset = 0
    for i, text in enumerate(sub_texts):
        est_page = max(1, int((char_offset / total_chars) * total_pages) + 1)
        chunks.append({
            "id": i + 1,
            "heading": f"Section {i + 1}",
            "level": 1,
            "text": text,
            "tokens": count_tokens(text),
            "start_page": est_page,
            "end_page": min(est_page + pages_per_chunk, total_pages),
        })
        char_offset += len(text)
    return chunks

# ---------------------------------------------------------------------------
# Core: extract + chunk a single file -> returns structured result
# ---------------------------------------------------------------------------
def extract_and_chunk(filepath: str, max_tokens: int, overlap: int) -> dict | None:
    """Extract text, detect boundaries, chunk. Returns result dict or None on failure."""
    filepath = os.path.abspath(filepath)
    ext = Path(filepath).suffix.lower()
    filename = Path(filepath).name

    if ext not in SUPPORTED_EXTENSIONS:
        print(f"Skipping unsupported file: {filename}", file=sys.stderr)
        return None

    print(f"Extracting text from {filename}...")
    blocks = extract_file(filepath)

    total_chars = sum(len(b.text) for b in blocks)
    if total_chars < 50:
        print(f"  WARNING: Very little text in {filename} ({total_chars} chars). Skipping.", file=sys.stderr)
        return None

    total_tokens = count_tokens(blocks_to_text(blocks))
    total_pages = max((b.page for b in blocks if b.page > 0), default=0)
    # txt/md files have no page concept; use 0
    print(f"  {total_chars} chars, ~{total_tokens} tokens" +
          (f", {total_pages} pages" if total_pages > 0 else ""))

    boundaries = detect_boundaries(blocks)
    high_conf = [b for b in boundaries if b.confidence >= 0.7]
    print(f"  {len(high_conf)} section boundaries detected")

    if len(high_conf) >= 3:
        chunking_mode = "structure_aware"
        chunks = chunk_by_structure(blocks, high_conf, max_tokens, overlap)
    else:
        chunking_mode = "token_based"
        chunks = chunk_by_tokens(blocks, max_tokens, overlap)

    print(f"  {len(chunks)} chunks ({chunking_mode})")

    doc_structure = []
    for b in sorted(boundaries, key=lambda x: x.index):
        if b.confidence >= 0.7:
            doc_structure.append({
                "heading": clean_heading(b.heading),
                "level": b.level,
                "page": b.page,
            })

    return {
        "source_file": filepath,
        "filename": filename,
        "file_type": ext.lstrip("."),
        "total_pages": total_pages,
        "total_tokens": total_tokens,
        "total_chars": total_chars,
        "chunking_mode": chunking_mode,
        "document_structure": doc_structure,
        "chunks": chunks,
    }

# ---------------------------------------------------------------------------
# Write output — single file mode
# ---------------------------------------------------------------------------
def write_single_file_output(result: dict, output_dir: str, max_tokens: int, overlap: int):
    """Write chunks and metadata for a single file."""
    chunks_dir = os.path.join(output_dir, "chunks")
    summaries_dir = os.path.join(output_dir, "summaries")
    os.makedirs(chunks_dir, exist_ok=True)
    os.makedirs(summaries_dir, exist_ok=True)

    chunk_metadata = []
    for chunk in result["chunks"]:
        chunk_file = f"chunk_{chunk['id']:03d}.txt"
        with open(os.path.join(chunks_dir, chunk_file), "w", encoding="utf-8") as f:
            f.write(chunk["text"])
        chunk_metadata.append({
            "id": chunk["id"],
            "file": f"chunks/{chunk_file}",
            "token_count": chunk["tokens"],
            "start_page": chunk["start_page"],
            "end_page": chunk["end_page"],
            "heading": chunk["heading"],
            "first_line": chunk["text"][:120].strip().replace("\n", " "),
        })

    metadata = {
        "mode": "single_file",
        "source_file": result["source_file"],
        "file_type": result["file_type"],
        "total_pages": result["total_pages"],
        "total_tokens": result["total_tokens"],
        "total_chars": result["total_chars"],
        "chunking_mode": result["chunking_mode"],
        "num_chunks": len(result["chunks"]),
        "overlap_tokens": overlap,
        "max_tokens_per_chunk": max_tokens,
        "document_structure": result["document_structure"],
        "chunks": chunk_metadata,
    }

    with open(os.path.join(output_dir, "metadata.json"), "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2)

    num = len(result["chunks"])
    print(f"\nOutput written to {output_dir}")
    print(f"  metadata.json + {num} chunks in chunks/")

    print(json.dumps({
        "status": "success",
        "mode": "single_file",
        "output_dir": output_dir,
        "num_chunks": num,
        "total_tokens": result["total_tokens"],
        "total_pages": result["total_pages"],
        "chunking_mode": result["chunking_mode"],
    }))

# ---------------------------------------------------------------------------
# Write output — multi-file (directory) mode
# ---------------------------------------------------------------------------
def write_multi_file_output(results: list[dict], source_dir: str,
                            output_dir: str, max_tokens: int, overlap: int):
    """Write chunks and metadata for multiple files from a directory."""
    chunks_dir = os.path.join(output_dir, "chunks")
    summaries_dir = os.path.join(output_dir, "summaries")
    os.makedirs(chunks_dir, exist_ok=True)
    os.makedirs(summaries_dir, exist_ok=True)

    files_metadata = []
    global_chunk_id = 0

    for file_idx, result in enumerate(results, start=1):
        prefix = f"f{file_idx:02d}"
        file_chunks = []

        for chunk in result["chunks"]:
            global_chunk_id += 1
            chunk_file = f"{prefix}_chunk_{chunk['id']:03d}.txt"
            with open(os.path.join(chunks_dir, chunk_file), "w", encoding="utf-8") as f:
                f.write(chunk["text"])
            file_chunks.append({
                "id": global_chunk_id,
                "local_id": chunk["id"],
                "file": f"chunks/{chunk_file}",
                "token_count": chunk["tokens"],
                "start_page": chunk["start_page"],
                "end_page": chunk["end_page"],
                "heading": chunk["heading"],
                "first_line": chunk["text"][:120].strip().replace("\n", " "),
            })

        files_metadata.append({
            "file_index": file_idx,
            "source_file": result["source_file"],
            "filename": result["filename"],
            "file_type": result["file_type"],
            "total_pages": result["total_pages"],
            "total_tokens": result["total_tokens"],
            "total_chars": result["total_chars"],
            "chunking_mode": result["chunking_mode"],
            "num_chunks": len(result["chunks"]),
            "document_structure": result["document_structure"],
            "chunks": file_chunks,
        })

    total_tokens = sum(r["total_tokens"] for r in results)
    total_chunks = global_chunk_id

    metadata = {
        "mode": "multi_file",
        "source_dir": os.path.abspath(source_dir),
        "num_files": len(results),
        "total_tokens": total_tokens,
        "total_chunks": total_chunks,
        "overlap_tokens": overlap,
        "max_tokens_per_chunk": max_tokens,
        "files": files_metadata,
    }

    with open(os.path.join(output_dir, "metadata.json"), "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2)

    print(f"\nOutput written to {output_dir}")
    print(f"  metadata.json: {len(results)} files, {total_chunks} total chunks")

    print(json.dumps({
        "status": "success",
        "mode": "multi_file",
        "output_dir": output_dir,
        "num_files": len(results),
        "total_chunks": total_chunks,
        "total_tokens": total_tokens,
        "files": [{"filename": r["filename"], "chunks": len(r["chunks"])} for r in results],
    }))

# ---------------------------------------------------------------------------
# Directory scanning
# ---------------------------------------------------------------------------
def find_supported_files(dirpath: str) -> list[str]:
    """Find all supported document files in a directory (non-recursive)."""
    files = []
    for entry in sorted(os.listdir(dirpath)):
        full = os.path.join(dirpath, entry)
        if os.path.isfile(full) and Path(entry).suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(full)
    return files

# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Chunk documents (PDF, DOCX, TXT, MD) for parallel LLM summarization. "
                    "Accepts a single file or a directory of files."
    )
    parser.add_argument("input_path", help="Path to a file or directory of files")
    parser.add_argument("output_dir", help="Directory for chunk output")
    parser.add_argument("--max-tokens", type=int, default=4000,
                        help="Target max tokens per chunk (default: 4000)")
    parser.add_argument("--overlap", type=int, default=200,
                        help="Token overlap between chunks (default: 200)")
    args = parser.parse_args()

    input_path = os.path.abspath(args.input_path)
    output_dir = os.path.abspath(args.output_dir)

    if os.path.isdir(input_path):
        # --- Directory mode ---
        files = find_supported_files(input_path)
        if not files:
            exts = ", ".join(sorted(SUPPORTED_EXTENSIONS))
            print(json.dumps({
                "error": f"No supported files found in {input_path}",
                "supported": exts,
            }), file=sys.stderr)
            sys.exit(1)

        print(f"Directory mode: found {len(files)} file(s) in {input_path}")
        for f in files:
            print(f"  - {Path(f).name}")
        print()

        results = []
        for filepath in files:
            result = extract_and_chunk(filepath, args.max_tokens, args.overlap)
            if result:
                results.append(result)

        if not results:
            print(json.dumps({"error": "No files could be processed"}), file=sys.stderr)
            sys.exit(1)

        write_multi_file_output(results, input_path, output_dir, args.max_tokens, args.overlap)

    elif os.path.isfile(input_path):
        # --- Single file mode ---
        ext = Path(input_path).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            exts = ", ".join(sorted(SUPPORTED_EXTENSIONS))
            print(json.dumps({
                "error": f"Unsupported file type: {ext}",
                "supported": exts,
            }), file=sys.stderr)
            sys.exit(1)

        result = extract_and_chunk(input_path, args.max_tokens, args.overlap)
        if not result:
            sys.exit(1)

        write_single_file_output(result, output_dir, args.max_tokens, args.overlap)

    else:
        print(json.dumps({"error": f"Path not found: {input_path}"}), file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
