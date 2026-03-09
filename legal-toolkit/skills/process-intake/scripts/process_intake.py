#!/usr/bin/env python3
"""
Legal Intake Processor

Takes raw client intake data (text notes, JSON forms, DOCX documents, CSV)
and structures it into standardized legal intake outputs: client profiles,
conflict check lists, document checklists, and statute of limitations warnings.

Outputs JSON to stdout for Claude to parse. Progress/errors go to stderr.
"""
import argparse
import csv
import json
import os
import re
import sys
from collections import defaultdict
from datetime import datetime, timedelta

import pandas as pd
try:
    import spacy
    HAS_SPACY = True
except ImportError:
    spacy = None
    HAS_SPACY = False
from dateutil import parser as dateutil_parser

# DOCX extraction
try:
    import docx as python_docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ---------------------------------------------------------------------------
# Utility helpers
# ---------------------------------------------------------------------------

def log(msg):
    """Print progress to stderr."""
    print(msg, file=sys.stderr)


# ---------------------------------------------------------------------------
# Input parsing
# ---------------------------------------------------------------------------

def parse_text_file(filepath):
    """Read plain text / markdown intake notes."""
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def parse_json_file(filepath):
    """Read structured JSON intake form data."""
    with open(filepath, "r", encoding="utf-8") as f:
        data = json.load(f)
    # Convert to text representation for NLP processing, but keep the structured data
    if isinstance(data, dict):
        text_parts = []
        for key, value in data.items():
            text_parts.append(f"{key}: {value}")
        return "\n".join(text_parts), data
    elif isinstance(data, list):
        # Multiple intake records
        all_text = []
        for item in data:
            if isinstance(item, dict):
                text_parts = [f"{k}: {v}" for k, v in item.items()]
                all_text.append("\n".join(text_parts))
        return "\n\n---\n\n".join(all_text), data
    return str(data), data


def parse_docx_file(filepath):
    """Read DOCX intake form."""
    if not HAS_DOCX:
        log("WARNING: python-docx not available. Cannot process DOCX files.")
        return ""
    doc = python_docx.Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def parse_csv_file(filepath):
    """Read CSV batch intake data. Returns list of row dicts."""
    rows = []
    with open(filepath, "r", encoding="utf-8-sig", errors="replace") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(dict(row))
    # Also create text representation
    text_parts = []
    for i, row in enumerate(rows):
        text_parts.append(f"--- Client {i+1} ---")
        for k, v in row.items():
            if v and v.strip():
                text_parts.append(f"{k}: {v}")
    return "\n".join(text_parts), rows


def load_input(filepath):
    """Load intake data from file. Returns (text, structured_data_or_None)."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in (".txt", ".md"):
        return parse_text_file(filepath), None
    elif ext == ".json":
        return parse_json_file(filepath)
    elif ext == ".docx":
        return parse_docx_file(filepath), None
    elif ext == ".csv":
        return parse_csv_file(filepath)
    else:
        log(f"ERROR: Unsupported file type: {ext}")
        sys.exit(1)


# ---------------------------------------------------------------------------
# Entity extraction
# ---------------------------------------------------------------------------

# Regex patterns for contact information
PHONE_PATTERN = re.compile(
    r"(?:(?:\+?1[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4})"
)
EMAIL_PATTERN = re.compile(
    r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
)
SSN_PATTERN = re.compile(r"\b\d{3}[-]?\d{2}[-]?\d{4}\b")
ADDRESS_PATTERN = re.compile(
    r"\d+\s+[\w\s]+(?:Street|St|Avenue|Ave|Boulevard|Blvd|Drive|Dr|Road|Rd|Lane|Ln|Way|Court|Ct|Circle|Cir|Place|Pl)\.?"
    r"(?:\s*(?:#|Apt|Suite|Ste|Unit)\s*\w+)?"
    r"(?:\s*,\s*\w+(?:\s+\w+)*\s*,\s*[A-Z]{2}\s*\d{5}(?:-\d{4})?)?",
    re.IGNORECASE,
)


def extract_contact_info(text):
    """Extract phone numbers, emails, and addresses from text."""
    phones = list(set(PHONE_PATTERN.findall(text)))
    emails = list(set(EMAIL_PATTERN.findall(text)))
    addresses = list(set(ADDRESS_PATTERN.findall(text)))

    # Clean up phone numbers
    phones = [re.sub(r"[^\d+()-.\s]", "", p).strip() for p in phones if len(re.sub(r"\D", "", p)) >= 7]

    return {
        "phones": phones[:5],  # Limit to avoid noise
        "emails": emails[:5],
        "addresses": [a.strip() for a in addresses[:3]],
    }


def chunk_text(text, max_length=100000):
    """Split text into chunks at paragraph boundaries."""
    if len(text) <= max_length:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        if start + max_length >= len(text):
            chunks.append(text[start:])
            break
        # Find last paragraph break before limit
        end = text.rfind('\n\n', start, start + max_length)
        if end == -1 or end <= start:
            # Fall back to sentence boundary
            end = text.rfind('. ', start, start + max_length)
        if end == -1 or end <= start:
            # Last resort: hard cut
            end = start + max_length
        else:
            end += 2  # Include the delimiter
        chunks.append(text[start:end])
        start = end
    return chunks


def extract_entities_nlp(nlp, text):
    """Extract named entities using spaCy."""
    # Process in chunks to handle large texts
    max_length = 100000
    entities = {
        "persons": [],
        "organizations": [],
        "dates": [],
        "money": [],
        "locations": [],
        "laws": [],
    }

    chunks = chunk_text(text, max_length)

    for chunk in chunks:
        try:
            doc = nlp(chunk)
        except Exception as e:
            log(f"  spaCy error: {e}")
            continue

        for ent in doc.ents:
            entry = {
                "text": ent.text.strip(),
                "context": ent.sent.text.strip()[:200] if ent.sent else "",
            }
            if ent.label_ == "PERSON":
                entities["persons"].append(entry)
            elif ent.label_ == "ORG":
                entities["organizations"].append(entry)
            elif ent.label_ == "DATE":
                entities["dates"].append(entry)
            elif ent.label_ == "MONEY":
                entities["money"].append(entry)
            elif ent.label_ in ("GPE", "LOC", "FAC"):
                entities["locations"].append(entry)
            elif ent.label_ == "LAW":
                entities["laws"].append(entry)

    return entities


# ---------------------------------------------------------------------------
# Matter type classification
# ---------------------------------------------------------------------------

MATTER_TYPE_KEYWORDS = {
    "personal_injury": [
        "accident", "injury", "injured", "crash", "collision", "slip", "fall",
        "medical", "hospital", "surgery", "disability", "pain", "suffering",
        "negligence", "liability", "damages", "insurance claim", "malpractice",
        "wrongful death", "product liability", "workers comp", "workers compensation",
    ],
    "family_law": [
        "divorce", "custody", "child support", "alimony", "separation",
        "prenuptial", "prenup", "visitation", "adoption", "guardianship",
        "domestic violence", "restraining order", "marriage", "spousal support",
        "marital", "paternity", "child protective",
    ],
    "criminal_defense": [
        "arrest", "criminal", "charge", "felony", "misdemeanor", "dui", "dwi",
        "assault", "battery", "theft", "drug", "possession", "probation",
        "parole", "sentencing", "plea", "defendant", "prosecutor", "bail",
        "indictment", "trial", "verdict", "conviction",
    ],
    "immigration": [
        "visa", "green card", "citizenship", "deportation", "asylum",
        "immigration", "naturalization", "uscis", "ice", "work permit",
        "i-130", "i-485", "h-1b", "daca", "refugee", "permanent resident",
        "travel document", "consulate", "embassy", "overstay",
    ],
    "corporate": [
        "corporation", "llc", "partnership", "shareholder", "merger",
        "acquisition", "formation", "bylaws", "operating agreement",
        "board of directors", "securities", "compliance", "contract",
        "non-compete", "trademark", "intellectual property",
    ],
    "real_estate": [
        "property", "deed", "mortgage", "lease", "landlord", "tenant",
        "eviction", "foreclosure", "title", "closing", "zoning",
        "real estate", "boundary", "easement", "hoa", "homeowner",
    ],
    "employment": [
        "employer", "employee", "termination", "fired", "discrimination",
        "harassment", "wage", "overtime", "wrongful termination", "retaliation",
        "eeoc", "ada", "fmla", "severance", "non-disclosure", "workplace",
    ],
    "estate_planning": [
        "will", "trust", "estate", "probate", "executor", "beneficiary",
        "power of attorney", "living will", "testament", "inheritance",
        "trustee", "heir", "decedent", "succession",
    ],
    "bankruptcy": [
        "bankruptcy", "chapter 7", "chapter 11", "chapter 13", "debt",
        "creditor", "debtor", "discharge", "filing", "insolvency",
        "means test", "automatic stay", "trustee", "liquidation",
    ],
    "intellectual_property": [
        "patent", "trademark", "copyright", "trade secret", "infringement",
        "ip", "licensing", "fair use", "dmca", "prior art",
        "registration", "intellectual property",
    ],
}


def classify_matter_type(text):
    """Classify the matter type based on keyword matching."""
    text_lower = text.lower()
    scores = {}

    for matter_type, keywords in MATTER_TYPE_KEYWORDS.items():
        score = 0
        matched_keywords = []
        for kw in keywords:
            count = text_lower.count(kw)
            if count > 0:
                score += count
                matched_keywords.append(kw)
        if score > 0:
            scores[matter_type] = {
                "score": score,
                "matched_keywords": matched_keywords,
            }

    if not scores:
        return "unknown", 0.0, []

    best_type = max(scores, key=lambda x: scores[x]["score"])
    best_score = scores[best_type]["score"]
    total_score = sum(s["score"] for s in scores.values())
    confidence = best_score / total_score if total_score > 0 else 0

    return best_type, round(confidence, 2), scores[best_type]["matched_keywords"]


# ---------------------------------------------------------------------------
# Conflict check preparation
# ---------------------------------------------------------------------------

def prepare_conflict_check(entities, text):
    """Extract and normalize all names for conflict checking."""
    conflicts = []
    seen = set()

    # Process persons
    for person in entities.get("persons", []):
        name = person["text"].strip()
        if len(name) < 2 or name.lower() in seen:
            continue
        seen.add(name.lower())

        # Generate variations
        variations = generate_name_variations(name)
        conflicts.append({
            "name": name,
            "type": "Person",
            "role": infer_role(name, text),
            "variations": variations,
        })

    # Process organizations
    for org in entities.get("organizations", []):
        name = org["text"].strip()
        if len(name) < 2 or name.lower() in seen:
            continue
        seen.add(name.lower())
        conflicts.append({
            "name": name,
            "type": "Organization",
            "role": infer_role(name, text),
            "variations": [name],
        })

    return conflicts


def generate_name_variations(name):
    """Generate variations of a person's name for conflict checking."""
    variations = [name]
    parts = name.split()

    if len(parts) >= 2:
        # Last name only
        variations.append(parts[-1])
        # First Last
        variations.append(f"{parts[0]} {parts[-1]}")
        # First initial + Last
        variations.append(f"{parts[0][0]}. {parts[-1]}")
        # Last, First
        variations.append(f"{parts[-1]}, {parts[0]}")

    if len(parts) >= 3:
        # First Middle Last -> First Last
        variations.append(f"{parts[0]} {parts[-1]}")
        # First M. Last
        variations.append(f"{parts[0]} {parts[1][0]}. {parts[-1]}")

    # Deduplicate while preserving order
    seen = set()
    unique = []
    for v in variations:
        if v.lower() not in seen:
            seen.add(v.lower())
            unique.append(v)
    return unique


def infer_role(name, text):
    """Attempt to infer the role of a named entity from context."""
    text_lower = text.lower()
    name_lower = name.lower()

    # Find sentences containing this name
    sentences = re.split(r"[.!?]\s+", text)
    relevant = [s.lower() for s in sentences if name_lower in s.lower()]

    role_keywords = {
        "Client": ["client", "caller", "contacted", "wants to", "needs", "seeking"],
        "Opposing Party": ["defendant", "opposing", "against", "sued", "vs", "versus"],
        "Spouse": ["spouse", "husband", "wife", "married to", "ex-husband", "ex-wife"],
        "Attorney": ["attorney", "lawyer", "counsel", "esq", "law firm", "representing"],
        "Witness": ["witness", "testified", "saw", "observed"],
        "Employer": ["employer", "company", "works for", "employed by"],
        "Insurance": ["insurance", "insurer", "adjuster", "carrier"],
        "Judge": ["judge", "honor", "court"],
        "Doctor": ["doctor", "dr.", "physician", "medical", "treated by"],
    }

    for role, keywords in role_keywords.items():
        for sent in relevant:
            for kw in keywords:
                if kw in sent:
                    return role

    return "Unknown"


# ---------------------------------------------------------------------------
# Document checklist generation
# ---------------------------------------------------------------------------

DOCUMENT_CHECKLISTS = {
    "personal_injury": {
        "critical": [
            "Medical records from all treating providers",
            "Police/incident report",
            "Insurance policy declarations page",
            "Photos of injuries and accident scene",
            "Emergency room records",
        ],
        "important": [
            "Medical bills and payment records",
            "Proof of lost wages (pay stubs, employer letter)",
            "Health insurance EOBs",
            "Prescription records",
            "Physical therapy records",
            "Vehicle repair estimates/photos (if auto accident)",
            "Witness statements or contact information",
        ],
        "supplemental": [
            "Prior medical records (pre-existing conditions)",
            "Surveillance footage",
            "Expert medical opinions",
            "Life care plan (if severe injury)",
            "Tax returns (for lost income verification)",
        ],
    },
    "family_law": {
        "critical": [
            "Marriage certificate",
            "Financial disclosure statements",
            "Most recent tax returns (3 years)",
            "Pay stubs (last 3 months)",
            "Bank statements (all accounts, last 6 months)",
        ],
        "important": [
            "Property deeds and mortgage statements",
            "Vehicle titles and loan statements",
            "Retirement/pension account statements",
            "Existing prenuptial or postnuptial agreements",
            "Custody or visitation agreements (if any)",
            "Children's birth certificates",
            "Health insurance cards",
        ],
        "supplemental": [
            "Investment account statements",
            "Business ownership documents",
            "Debt statements (credit cards, loans)",
            "Social media evidence",
            "Communication records (texts, emails)",
            "School records for children",
        ],
    },
    "criminal_defense": {
        "critical": [
            "Arrest report and charging documents",
            "Bail/bond information",
            "Prior criminal history (if known)",
            "Police body camera or dashcam footage",
            "Witness statements from police report",
        ],
        "important": [
            "Employment verification",
            "Character reference letters",
            "Alibi evidence",
            "Medical records (if injury involved)",
            "Communication records relevant to charges",
            "Financial records (if financial crime)",
        ],
        "supplemental": [
            "Rehabilitation records",
            "Community involvement documentation",
            "Military service records",
            "Educational records",
        ],
    },
    "immigration": {
        "critical": [
            "Valid passport (all pages)",
            "Current visa or immigration status documentation",
            "I-94 arrival/departure record",
            "Birth certificate with translation",
            "Employment authorization document (if applicable)",
        ],
        "important": [
            "Prior immigration filings and receipts",
            "Employment verification letter",
            "Pay stubs (last 3 months)",
            "Tax returns (last 3 years)",
            "Marriage certificate (if family-based)",
            "Spouse's citizenship/residency proof",
        ],
        "supplemental": [
            "Educational diplomas and transcripts",
            "Professional certifications",
            "Photographs (passport-style)",
            "Affidavits of support",
            "Evidence of residence in the US",
            "Police clearance certificates",
        ],
    },
    "corporate": {
        "critical": [
            "Articles of incorporation / formation documents",
            "Operating agreement or bylaws",
            "Shareholder/member agreements",
            "Current contracts relevant to the matter",
            "Board minutes and resolutions",
        ],
        "important": [
            "Financial statements (last 3 years)",
            "Tax returns (business, last 3 years)",
            "Insurance policies (general liability, D&O)",
            "Employment agreements",
            "Non-compete/NDA agreements",
        ],
        "supplemental": [
            "Intellectual property registrations",
            "Regulatory filings and licenses",
            "Prior legal opinions or memoranda",
            "Correspondence related to the matter",
        ],
    },
    "real_estate": {
        "critical": [
            "Property deed",
            "Purchase/sale agreement",
            "Title report or title insurance policy",
            "Mortgage documents",
            "Property survey",
        ],
        "important": [
            "Lease agreements (if rental property)",
            "HOA documents and bylaws",
            "Property tax statements",
            "Insurance policy",
            "Inspection reports",
            "Zoning and permit records",
        ],
        "supplemental": [
            "Environmental reports",
            "Appraisal reports",
            "Correspondence with other parties",
            "Photos of property",
        ],
    },
    "employment": {
        "critical": [
            "Employment agreement/offer letter",
            "Termination letter or documentation",
            "Pay stubs (last 12 months)",
            "Employee handbook/policies",
            "Performance reviews",
        ],
        "important": [
            "Written warnings or disciplinary records",
            "Emails/communications about the dispute",
            "EEOC charge or complaint (if filed)",
            "Benefits documentation",
            "Non-compete or NDA agreements",
            "Company org chart",
        ],
        "supplemental": [
            "Witness contact information",
            "Severance agreement (if offered)",
            "Social media evidence",
            "Job postings (for replacement evidence)",
            "Comparable salary data",
        ],
    },
    "estate_planning": {
        "critical": [
            "Existing will or trust documents",
            "List of all assets (real property, accounts, investments)",
            "List of all debts and liabilities",
            "Beneficiary designations (life insurance, retirement accounts)",
            "Government-issued ID",
        ],
        "important": [
            "Existing power of attorney documents",
            "Healthcare directive / living will",
            "Deed to real property",
            "Recent tax returns",
            "Life insurance policies",
            "Retirement account statements",
        ],
        "supplemental": [
            "Business ownership documents",
            "Prenuptial/postnuptial agreements",
            "Divorce decrees",
            "Special needs trust documentation (if applicable)",
            "Charitable giving intentions",
        ],
    },
    "bankruptcy": {
        "critical": [
            "List of all creditors with amounts owed",
            "Bank statements (all accounts, last 6 months)",
            "Tax returns (last 2 years)",
            "Pay stubs (last 6 months)",
            "List of all assets and property",
        ],
        "important": [
            "Mortgage statements",
            "Vehicle loan statements",
            "Credit card statements (last 6 months)",
            "Collection letters and lawsuits",
            "Proof of expenses (utilities, rent, etc.)",
            "Previous bankruptcy filings (if any)",
        ],
        "supplemental": [
            "Loan modifications or hardship letters",
            "Property appraisals",
            "Medical bills contributing to debt",
            "Retirement account statements",
        ],
    },
    "intellectual_property": {
        "critical": [
            "Registration certificates (patent, trademark, copyright)",
            "Application filings and correspondence",
            "Evidence of creation/invention (dates, records)",
            "Evidence of infringement",
            "Licensing agreements",
        ],
        "important": [
            "Prior art search results (patents)",
            "Market evidence (for damages)",
            "Correspondence with infringing party",
            "Business plans showing IP value",
            "Domain registration records",
        ],
        "supplemental": [
            "Expert opinions or analyses",
            "Industry standards and practices",
            "International filings",
            "Social media evidence of infringement",
        ],
    },
}


def generate_document_checklist(matter_type):
    """Generate a matter-type-specific document checklist."""
    checklist = DOCUMENT_CHECKLISTS.get(matter_type)
    if not checklist:
        return {
            "matter_type": matter_type,
            "note": "No specific checklist available for this matter type. Request general documents: ID, relevant contracts, correspondence, and financial records.",
            "critical": ["Government-issued identification", "All relevant contracts and agreements", "Correspondence related to the matter"],
            "important": ["Financial records", "Timeline of events", "Contact information for relevant parties"],
            "supplemental": [],
        }
    return {
        "matter_type": matter_type,
        "critical": checklist.get("critical", []),
        "important": checklist.get("important", []),
        "supplemental": checklist.get("supplemental", []),
    }


# ---------------------------------------------------------------------------
# Statute of limitations
# ---------------------------------------------------------------------------

# Simplified SOL data (years) by matter type and jurisdiction
# In practice, this would be a comprehensive database
SOL_DATA = {
    "personal_injury": {
        "default": 2,
        "CA": 2, "NY": 3, "TX": 2, "FL": 4, "IL": 2,
        "PA": 2, "OH": 2, "GA": 2, "NC": 3, "MI": 3,
        "NJ": 2, "VA": 2, "WA": 3, "MA": 3, "AZ": 2,
        "CO": 2, "MD": 3, "MN": 6, "WI": 3, "MO": 5,
    },
    "family_law": {
        "note": "Family law matters generally do not have a statute of limitations for filing. However, specific claims within family law (e.g., property division after divorce) may have deadlines.",
    },
    "criminal_defense": {
        "note": "Criminal statutes of limitations vary widely by offense type and jurisdiction. Consult jurisdiction-specific criminal code.",
    },
    "employment": {
        "default": 1,  # EEOC charge filing (general)
        "CA": 3, "NY": 3, "TX": 1, "FL": 1, "IL": 1,
        "note": "EEOC charges must typically be filed within 180-300 days. State deadlines vary.",
    },
    "real_estate": {
        "default": 4,
        "CA": 4, "NY": 6, "TX": 4, "FL": 5, "IL": 5,
    },
    "corporate": {
        "default": 4,
        "CA": 4, "NY": 6, "TX": 4, "FL": 5,
        "note": "Depends heavily on the specific claim (breach of contract, fraud, etc.)",
    },
    "bankruptcy": {
        "note": "Bankruptcy can generally be filed at any time. Chapter 7 has an 8-year refiling limitation. Chapter 13 has varying limitations.",
    },
    "intellectual_property": {
        "default": 3,
        "note": "Patent: 6 years for damages. Trademark: varies. Copyright: 3 years from discovery.",
    },
    "estate_planning": {
        "note": "Estate planning itself has no statute of limitations. Probate contests typically have short deadlines (often 120 days from admission).",
    },
    "immigration": {
        "note": "Immigration matters have specific filing deadlines tied to visa categories, not general statutes of limitations.",
    },
}


def calculate_sol_warnings(matter_type, jurisdiction, dates_found):
    """Calculate statute of limitations warnings."""
    warnings = []
    sol_info = SOL_DATA.get(matter_type, {})

    # If there's just a note (no numeric SOL)
    if "note" in sol_info and "default" not in sol_info:
        return [{
            "matter_type": matter_type,
            "jurisdiction": jurisdiction or "unspecified",
            "note": sol_info["note"],
            "warning_level": "info",
        }]

    # Get SOL years
    sol_years = sol_info.get(jurisdiction, sol_info.get("default"))
    if sol_years is None:
        return [{
            "matter_type": matter_type,
            "jurisdiction": jurisdiction or "unspecified",
            "note": "Could not determine statute of limitations. Consult jurisdiction-specific rules.",
            "warning_level": "info",
        }]

    # Try to find an incident date from extracted dates
    incident_date = None
    for date_entry in dates_found:
        date_text = date_entry.get("text", "")
        context = date_entry.get("context", "").lower()
        # Look for dates associated with incident-related keywords
        incident_keywords = ["accident", "incident", "injury", "occurred", "happened", "event", "arrest", "terminated", "fired"]
        if any(kw in context for kw in incident_keywords):
            try:
                incident_date = dateutil_parser.parse(date_text, fuzzy=True)
                break
            except Exception:
                continue

    if incident_date is None:
        # Try the first date that looks like it could be past
        now = datetime.now()
        for date_entry in dates_found:
            try:
                parsed = dateutil_parser.parse(date_entry["text"], fuzzy=True)
                if parsed < now:
                    incident_date = parsed
                    break
            except Exception:
                continue

    if incident_date:
        sol_deadline = incident_date + timedelta(days=sol_years * 365)
        days_remaining = (sol_deadline - datetime.now()).days

        if days_remaining < 0:
            warning_level = "expired"
        elif days_remaining < 30:
            warning_level = "urgent"
        elif days_remaining < 90:
            warning_level = "caution"
        else:
            warning_level = "normal"

        warnings.append({
            "matter_type": matter_type,
            "jurisdiction": jurisdiction or "unspecified",
            "sol_years": sol_years,
            "incident_date": incident_date.strftime("%Y-%m-%d"),
            "sol_deadline": sol_deadline.strftime("%Y-%m-%d"),
            "days_remaining": max(0, days_remaining),
            "warning_level": warning_level,
        })
    else:
        warnings.append({
            "matter_type": matter_type,
            "jurisdiction": jurisdiction or "unspecified",
            "sol_years": sol_years,
            "note": f"SOL is {sol_years} years from incident date, but no incident date was detected. Please verify manually.",
            "warning_level": "info",
        })

    if "note" in sol_info:
        warnings.append({
            "additional_note": sol_info["note"],
            "warning_level": "info",
        })

    return warnings


# ---------------------------------------------------------------------------
# Output generation
# ---------------------------------------------------------------------------

def write_outputs(client_profile, conflict_list, checklist, sol_warnings,
                  intake_text, output_dir):
    """Write all output files."""

    # 1. client_profile.json
    log("  Writing client_profile.json...")
    profile_path = os.path.join(output_dir, "client_profile.json")
    with open(profile_path, "w", encoding="utf-8") as f:
        json.dump(client_profile, f, indent=2, default=str)

    # 2. conflict_check.xlsx
    log("  Writing conflict_check.xlsx...")
    if conflict_list:
        rows = []
        for entry in conflict_list:
            rows.append({
                "Name": entry["name"],
                "Type": entry["type"],
                "Role": entry["role"],
                "Variations": "; ".join(entry["variations"]),
            })
        df = pd.DataFrame(rows)
        df.to_excel(os.path.join(output_dir, "conflict_check.xlsx"),
                     index=False, engine="xlsxwriter")

    # 3. document_checklist.json
    log("  Writing document_checklist.json...")
    checklist_path = os.path.join(output_dir, "document_checklist.json")
    with open(checklist_path, "w", encoding="utf-8") as f:
        json.dump(checklist, f, indent=2)

    # 4. intake_summary.txt
    log("  Writing intake_summary.txt...")
    summary_lines = [
        "=" * 60,
        "LEGAL INTAKE PROCESSING SUMMARY",
        "=" * 60,
        "",
    ]

    # Client info
    summary_lines.append("CLIENT INFORMATION:")
    if client_profile.get("client_name"):
        summary_lines.append(f"  Name: {client_profile['client_name']}")
    contact = client_profile.get("contact_info", {})
    if contact.get("phones"):
        summary_lines.append(f"  Phone: {', '.join(contact['phones'])}")
    if contact.get("emails"):
        summary_lines.append(f"  Email: {', '.join(contact['emails'])}")
    if contact.get("addresses"):
        summary_lines.append(f"  Address: {contact['addresses'][0]}")

    summary_lines.append("")
    summary_lines.append("MATTER CLASSIFICATION:")
    summary_lines.append(f"  Type: {client_profile.get('matter_type', 'unknown')}")
    summary_lines.append(f"  Confidence: {client_profile.get('matter_confidence', 0):.0%}")
    if client_profile.get("matter_keywords"):
        summary_lines.append(f"  Keywords: {', '.join(client_profile['matter_keywords'][:5])}")

    summary_lines.append("")
    summary_lines.append("KEY DATES:")
    for date_entry in client_profile.get("dates", [])[:5]:
        summary_lines.append(f"  {date_entry.get('text', '')}: {date_entry.get('context', '')[:80]}")

    summary_lines.append("")
    summary_lines.append("MONETARY AMOUNTS:")
    for money_entry in client_profile.get("money", [])[:5]:
        summary_lines.append(f"  {money_entry.get('text', '')}: {money_entry.get('context', '')[:80]}")

    summary_lines.append("")
    summary_lines.append(f"CONFLICT CHECK ENTITIES: {len(conflict_list)}")
    for entry in conflict_list[:10]:
        summary_lines.append(f"  [{entry['type']}] {entry['name']} - Role: {entry['role']}")

    summary_lines.append("")
    summary_lines.append("DOCUMENT CHECKLIST:")
    summary_lines.append(f"  Matter type: {checklist.get('matter_type', 'unknown')}")
    summary_lines.append(f"  Critical documents: {len(checklist.get('critical', []))}")
    summary_lines.append(f"  Important documents: {len(checklist.get('important', []))}")
    summary_lines.append(f"  Supplemental documents: {len(checklist.get('supplemental', []))}")

    if sol_warnings:
        summary_lines.append("")
        summary_lines.append("STATUTE OF LIMITATIONS:")
        for w in sol_warnings:
            if w.get("sol_deadline"):
                summary_lines.append(
                    f"  Deadline: {w['sol_deadline']} "
                    f"({w.get('days_remaining', '?')} days remaining) "
                    f"[{w.get('warning_level', '').upper()}]"
                )
            if w.get("note"):
                summary_lines.append(f"  Note: {w['note']}")
            if w.get("additional_note"):
                summary_lines.append(f"  Note: {w['additional_note']}")

    summary_lines.extend([
        "",
        "=" * 60,
        "Output files:",
        "  client_profile.json - Structured client data",
        "  conflict_check.xlsx - Entities for conflict checking",
        "  document_checklist.json - Required documents list",
        "  intake_summary.txt - This summary",
    ])
    if sol_warnings:
        summary_lines.append("  sol_warning.txt - Statute of limitations warnings")
    summary_lines.append("=" * 60)

    summary_text = "\n".join(summary_lines)
    with open(os.path.join(output_dir, "intake_summary.txt"), "w", encoding="utf-8") as f:
        f.write(summary_text)

    # 5. sol_warning.txt
    if sol_warnings:
        log("  Writing sol_warning.txt...")
        sol_lines = [
            "STATUTE OF LIMITATIONS WARNINGS",
            "=" * 40,
            "",
            "DISCLAIMER: These are estimated deadlines based on general SOL rules.",
            "Always verify with current statutes for your specific jurisdiction and claim type.",
            "",
        ]
        for w in sol_warnings:
            if w.get("sol_deadline"):
                level = w.get("warning_level", "").upper()
                prefix = ""
                if level == "EXPIRED":
                    prefix = "[EXPIRED] "
                elif level == "URGENT":
                    prefix = "[URGENT] "
                elif level == "CAUTION":
                    prefix = "[CAUTION] "

                sol_lines.append(f"{prefix}Matter: {w.get('matter_type', '')}")
                sol_lines.append(f"  Jurisdiction: {w.get('jurisdiction', 'unspecified')}")
                sol_lines.append(f"  SOL Period: {w.get('sol_years', '?')} years")
                sol_lines.append(f"  Incident Date: {w.get('incident_date', 'unknown')}")
                sol_lines.append(f"  Deadline: {w['sol_deadline']}")
                sol_lines.append(f"  Days Remaining: {w.get('days_remaining', '?')}")
                sol_lines.append("")
            if w.get("note"):
                sol_lines.append(f"Note: {w['note']}")
                sol_lines.append("")
            if w.get("additional_note"):
                sol_lines.append(f"Additional: {w['additional_note']}")
                sol_lines.append("")

        with open(os.path.join(output_dir, "sol_warning.txt"), "w", encoding="utf-8") as f:
            f.write("\n".join(sol_lines))

    return {
        "client_name": client_profile.get("client_name", ""),
        "matter_type": client_profile.get("matter_type", "unknown"),
        "matter_confidence": client_profile.get("matter_confidence", 0),
        "conflict_check_entities": len(conflict_list),
        "dates_found": len(client_profile.get("dates", [])),
        "money_found": len(client_profile.get("money", [])),
        "sol_warnings": len(sol_warnings),
        "sol_warning_level": sol_warnings[0].get("warning_level", "info") if sol_warnings else "none",
        "document_checklist_items": (
            len(checklist.get("critical", [])) +
            len(checklist.get("important", [])) +
            len(checklist.get("supplemental", []))
        ),
        "output_dir": output_dir,
        "files_generated": [
            "client_profile.json",
            "conflict_check.xlsx",
            "document_checklist.json",
            "intake_summary.txt",
        ] + (["sol_warning.txt"] if sol_warnings else []),
    }


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Legal Intake Processor")
    parser.add_argument("--input", required=True, help="Intake file path")
    parser.add_argument("--output-dir", required=True, help="Output directory")
    parser.add_argument("--matter-type", default="auto",
                        help="Matter type (auto to detect, or specify)")
    parser.add_argument("--jurisdiction", default=None,
                        help="Jurisdiction code (e.g., CA, NY, TX)")
    args = parser.parse_args()

    input_path = os.path.abspath(args.input)
    output_dir = os.path.abspath(args.output_dir)
    os.makedirs(output_dir, exist_ok=True)

    if not os.path.exists(input_path):
        log(f"ERROR: Input file does not exist: {input_path}")
        sys.exit(1)

    # Load spaCy model
    if not HAS_SPACY:
        log("ERROR: spaCy is not installed. Run check_dependencies.py first.")
        result = {"status": "error", "error": "spaCy is not installed. Run check_dependencies.py first."}
        print(json.dumps(result))
        sys.exit(1)
    log("Loading spaCy model: en_core_web_sm")
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        log("Model not found. Attempting download...")
        import subprocess
        subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"],
                       capture_output=True, text=True)
        try:
            nlp = spacy.load("en_core_web_sm")
        except OSError:
            log("ERROR: Could not load spaCy model.")
            sys.exit(1)

    nlp.max_length = 2_000_000

    # Load input
    log(f"Loading intake data: {input_path}")
    text, structured_data = load_input(input_path)

    if not text or not text.strip():
        log("ERROR: No content could be extracted from the input file.")
        sys.exit(1)

    log(f"  Extracted {len(text)} characters")

    # Extract contact info using regex
    log("Extracting contact information...")
    contact_info = extract_contact_info(text)

    # Extract entities using NLP
    log("Running named entity recognition...")
    entities = extract_entities_nlp(nlp, text)
    log(f"  Persons: {len(entities['persons'])}")
    log(f"  Organizations: {len(entities['organizations'])}")
    log(f"  Dates: {len(entities['dates'])}")
    log(f"  Money: {len(entities['money'])}")
    log(f"  Locations: {len(entities['locations'])}")

    # Classify matter type
    if args.matter_type == "auto":
        log("Classifying matter type...")
        matter_type, confidence, keywords = classify_matter_type(text)
        log(f"  Detected: {matter_type} (confidence: {confidence:.0%})")
    else:
        matter_type = args.matter_type
        confidence = 1.0
        keywords = []
        log(f"  Matter type (provided): {matter_type}")

    # Determine client name (first person entity, or from structured data)
    client_name = ""
    if structured_data and isinstance(structured_data, dict):
        # Try common field names
        for field in ["client_name", "name", "client", "full_name", "applicant"]:
            if field in structured_data and structured_data[field]:
                client_name = str(structured_data[field])
                break
    if not client_name and entities["persons"]:
        # Use the first person mentioned (likely the client in intake notes)
        client_name = entities["persons"][0]["text"]

    # Build client profile
    client_profile = {
        "client_name": client_name,
        "contact_info": contact_info,
        "matter_type": matter_type,
        "matter_confidence": confidence,
        "matter_keywords": keywords,
        "persons": [{"text": p["text"], "context": p["context"]} for p in entities["persons"][:20]],
        "organizations": [{"text": o["text"], "context": o["context"]} for o in entities["organizations"][:20]],
        "dates": [{"text": d["text"], "context": d["context"]} for d in entities["dates"][:20]],
        "money": [{"text": m["text"], "context": m["context"]} for m in entities["money"][:20]],
        "locations": [{"text": l["text"], "context": l["context"]} for l in entities["locations"][:20]],
        "opposing_parties": [],
        "source_file": os.path.basename(input_path),
    }

    # Try to identify opposing parties
    text_lower = text.lower()
    opposing_keywords = ["against", "defendant", "opposing", "respondent", "vs", "versus"]
    for person in entities["persons"]:
        context_lower = person.get("context", "").lower()
        if any(kw in context_lower for kw in opposing_keywords):
            if person["text"] != client_name:
                client_profile["opposing_parties"].append(person["text"])

    # If structured data has opposing party fields
    if structured_data and isinstance(structured_data, dict):
        for field in ["opposing_party", "defendant", "respondent", "adverse_party"]:
            if field in structured_data and structured_data[field]:
                client_profile["opposing_parties"].append(str(structured_data[field]))

    client_profile["opposing_parties"] = list(set(client_profile["opposing_parties"]))

    # Prepare conflict check
    log("Preparing conflict check list...")
    conflict_list = prepare_conflict_check(entities, text)
    log(f"  Entities for conflict check: {len(conflict_list)}")

    # Generate document checklist
    log("Generating document checklist...")
    checklist = generate_document_checklist(matter_type)

    # Calculate SOL warnings
    sol_warnings = []
    if args.jurisdiction:
        log(f"Calculating statute of limitations for {args.jurisdiction}...")
        sol_warnings = calculate_sol_warnings(
            matter_type, args.jurisdiction, entities.get("dates", [])
        )
    else:
        log("No jurisdiction provided; skipping SOL calculation.")

    # Write all outputs
    log("\nWriting outputs...")
    result = write_outputs(
        client_profile, conflict_list, checklist, sol_warnings,
        text, output_dir
    )

    log(f"\nIntake processing complete. Output: {output_dir}")

    # Print JSON to stdout for Claude
    print(json.dumps(result, indent=2, default=str))


if __name__ == "__main__":
    main()
