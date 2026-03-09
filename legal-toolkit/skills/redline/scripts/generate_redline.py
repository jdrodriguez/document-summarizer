#!/usr/bin/env python3
"""
Contract redline generator for the legal-redline skill.

Takes two versions of a contract (.docx) and generates a tracked-changes
redline document with native Word revision markup, plus risk-rated
change analysis.

Usage:
    python3 generate_redline.py --original <path.docx> --revised <path.docx> \
        --output-dir <dir>
"""
import argparse
import datetime
import difflib
import json
import os
import re
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Dependency imports
# ---------------------------------------------------------------------------
try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    DocxDocument = None
    print("ERROR: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(2)


# Word OpenXML namespace
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"

# Revision author and date
AUTHOR = "Legal Redline Generator"
REV_DATE = datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")


# ---------------------------------------------------------------------------
# Risk classification keywords
# ---------------------------------------------------------------------------
HIGH_RISK_KEYWORDS = [
    # Material terms
    r'\b(?:price|pricing|cost|fee|payment|compensation|salary|wage)\b',
    r'\b(?:liability|liable|damages|penalty|penalties|liquidated\s+damages)\b',
    r'\b(?:indemnif|indemnity|hold\s+harmless)\b',
    r'\b(?:termination|terminate|cancel|cancellation)\b',
    r'\b(?:governing\s+law|jurisdiction|arbitration|dispute\s+resolution|venue)\b',
    r'\b(?:warranty|warranties|guarantee|representation)\b',
    r'\b(?:intellectual\s+property|IP\s+rights|patent|copyright|trademark)\b',
    r'\b(?:confidential|non-disclosure|NDA|trade\s+secret)\b',
    r'\b(?:limitation\s+of\s+liability|cap|aggregate|maximum)\b',
    r'\b(?:non-compete|non-solicitation|restrictive\s+covenant)\b',
    r'\b(?:insurance|coverage|deductible)\b',
    r'\b(?:assignment|transfer|subcontract)\b',
    r'\b(?:force\s+majeure)\b',
]

MEDIUM_RISK_KEYWORDS = [
    r'\b(?:obligation|duty|shall|must|required|responsible)\b',
    r'\b(?:condition|contingent|subject\s+to|provided\s+that)\b',
    r'\b(?:notice|notification|written\s+notice|days?\s+(?:prior|advance))\b',
    r'\b(?:renewal|extension|option|right\s+of\s+first)\b',
    r'\b(?:deliverable|milestone|deadline|timeline|schedule)\b',
    r'\b(?:scope|specification|standard|requirement)\b',
    r'\b(?:approve|approval|consent)\b',
    r'\b(?:exclusiv|sole|preferr)\b',
    r'\b(?:audit|inspection|review\s+right)\b',
]


def classify_change(old_text: str, new_text: str) -> dict:
    """
    Classify a change by category and risk level.
    Returns: {"category": str, "risk": str, "risk_reasons": [...]}
    """
    combined = (old_text + " " + new_text).lower()

    # Check for high-risk keywords
    high_matches = []
    for pattern in HIGH_RISK_KEYWORDS:
        if re.search(pattern, combined, re.IGNORECASE):
            high_matches.append(pattern.replace(r'\b', '').replace('(?:', '').replace(')', '').replace('|', '/'))

    # Check for medium-risk keywords
    medium_matches = []
    for pattern in MEDIUM_RISK_KEYWORDS:
        if re.search(pattern, combined, re.IGNORECASE):
            medium_matches.append(pattern.replace(r'\b', '').replace('(?:', '').replace(')', '').replace('|', '/'))

    # Determine risk level
    if high_matches:
        risk = "HIGH"
        reasons = [f"Contains material terms: {', '.join(high_matches[:3])}"]
    elif medium_matches:
        risk = "MEDIUM"
        reasons = [f"Contains contractual terms: {', '.join(medium_matches[:3])}"]
    else:
        risk = "LOW"
        reasons = ["Administrative or formatting change"]

    # Determine category
    if high_matches:
        if any(re.search(r'(?:liability|indemnif|damages|penalty|limitation)', combined, re.I)):
            category = "Risk-relevant"
        else:
            category = "Substantive"
    elif medium_matches:
        category = "Substantive"
    else:
        # Check if it's just minor wording or formatting
        old_words = set(old_text.lower().split())
        new_words = set(new_text.lower().split())
        if len(old_words.symmetric_difference(new_words)) <= 3:
            category = "Administrative"
        else:
            category = "Substantive"

    return {
        "category": category,
        "risk": risk,
        "risk_reasons": reasons,
    }


# ---------------------------------------------------------------------------
# Text extraction from DOCX
# ---------------------------------------------------------------------------
def extract_paragraphs(doc: "DocxDocument") -> list[dict]:
    """Extract paragraphs with their styles and text."""
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else "Normal"
        paragraphs.append({
            "index": i,
            "text": text,
            "style": style,
            "runs": [{"text": run.text, "bold": run.bold, "italic": run.italic}
                     for run in para.runs],
        })
    return paragraphs


# ---------------------------------------------------------------------------
# Section detection for contracts
# ---------------------------------------------------------------------------
SECTION_RE = re.compile(
    r'^(?:(?:ARTICLE|SECTION|CHAPTER|PART)\s+[\dIVXLCDM]+|'
    r'(?:\d+\.)+\s+|'
    r'[IVXLCDM]+\.\s+)',
    re.IGNORECASE
)


def detect_section(text: str) -> str:
    """Try to identify which contract section a paragraph belongs to."""
    if SECTION_RE.match(text):
        return text[:80].strip()
    if len(text) < 100 and text == text.upper() and len(text.split()) >= 2:
        return text[:80].strip()
    return ""


def assign_sections(paragraphs: list[dict]) -> list[dict]:
    """Assign each paragraph to a contract section."""
    current_section = "Preamble"
    for para in paragraphs:
        section = detect_section(para["text"])
        if section:
            current_section = section
        para["section"] = current_section
    return paragraphs


# ---------------------------------------------------------------------------
# Tracked changes document generation using OpenXML
# ---------------------------------------------------------------------------
def create_revision_id():
    """Generate a unique revision ID."""
    create_revision_id.counter = getattr(create_revision_id, 'counter', 0) + 1
    return str(create_revision_id.counter)


def make_run_properties(bold=False, italic=False, color=None, strikethrough=False, underline=False):
    """Create a w:rPr element with formatting."""
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if strikethrough:
        strike = OxmlElement('w:strike')
        rPr.append(strike)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    return rPr


def create_run_element(text: str, rPr=None):
    """Create a w:r element with text."""
    r = OxmlElement('w:r')
    if rPr is not None:
        r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def create_insertion_element(text: str):
    """Create a w:ins element (tracked insertion) wrapping a run."""
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), create_revision_id())
    ins.set(qn('w:author'), AUTHOR)
    ins.set(qn('w:date'), REV_DATE)

    rPr = make_run_properties(color='FF0000', underline=True)
    r = create_run_element(text, rPr)
    ins.append(r)
    return ins


def create_deletion_element(text: str):
    """Create a w:del element (tracked deletion) wrapping a delText run."""
    delete = OxmlElement('w:del')
    delete.set(qn('w:id'), create_revision_id())
    delete.set(qn('w:author'), AUTHOR)
    delete.set(qn('w:date'), REV_DATE)

    r = OxmlElement('w:r')
    rPr = make_run_properties(color='FF0000', strikethrough=True)
    r.append(rPr)

    delText = OxmlElement('w:delText')
    delText.set(qn('xml:space'), 'preserve')
    delText.text = text
    r.append(delText)

    delete.append(r)
    return delete


def build_tracked_paragraph(old_text: str, new_text: str, style_name: str = "Normal"):
    """
    Build a paragraph element with tracked changes markup.
    Uses word-level diff to create fine-grained insertions and deletions.
    """
    p = OxmlElement('w:p')

    # Add paragraph properties for style
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name.replace(" ", ""))
    pPr.append(pStyle)
    p.append(pPr)

    if not old_text and new_text:
        # Entire paragraph is an insertion
        ins = create_insertion_element(new_text)
        p.append(ins)
        return p

    if old_text and not new_text:
        # Entire paragraph is a deletion
        delete = create_deletion_element(old_text)
        p.append(delete)
        return p

    # Word-level diff for modifications
    old_words = old_text.split()
    new_words = new_text.split()
    matcher = difflib.SequenceMatcher(None, old_words, new_words)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            text = " ".join(old_words[i1:i2])
            # Add space before if not first element
            if list(p) and list(p)[-1].tag != qn('w:pPr'):
                text = " " + text
            r = create_run_element(text)
            p.append(r)

        elif tag == "replace":
            old_chunk = " ".join(old_words[i1:i2])
            new_chunk = " ".join(new_words[j1:j2])
            # Add space before
            if list(p) and list(p)[-1].tag != qn('w:pPr'):
                old_chunk = " " + old_chunk
                new_chunk = " " + new_chunk
            delete = create_deletion_element(old_chunk)
            ins = create_insertion_element(new_chunk)
            p.append(delete)
            p.append(ins)

        elif tag == "delete":
            old_chunk = " ".join(old_words[i1:i2])
            if list(p) and list(p)[-1].tag != qn('w:pPr'):
                old_chunk = " " + old_chunk
            delete = create_deletion_element(old_chunk)
            p.append(delete)

        elif tag == "insert":
            new_chunk = " ".join(new_words[j1:j2])
            if list(p) and list(p)[-1].tag != qn('w:pPr'):
                new_chunk = " " + new_chunk
            ins = create_insertion_element(new_chunk)
            p.append(ins)

    return p


def generate_redline_docx(orig_paras: list[dict], rev_paras: list[dict],
                          output_path: str) -> list[dict]:
    """
    Generate a .docx file with tracked changes markup.
    Returns list of change records.
    """
    doc = DocxDocument()

    # Set up document defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("REDLINE COMPARISON")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # Spacer

    # Legend
    legend = doc.add_paragraph()
    legend_run = legend.add_run("Legend: ")
    legend_run.bold = True
    ins_run = legend.add_run("Inserted text")
    ins_run.font.color.rgb = RGBColor(255, 0, 0)
    ins_run.underline = True
    legend.add_run(" | ")
    del_run = legend.add_run("Deleted text")
    del_run.font.color.rgb = RGBColor(255, 0, 0)
    del_run.font.strike = True

    doc.add_paragraph()  # Spacer

    # Perform paragraph-level alignment
    orig_texts = [p["text"] for p in orig_paras]
    rev_texts = [p["text"] for p in rev_paras]
    matcher = difflib.SequenceMatcher(None, orig_texts, rev_texts)

    changes = []
    change_id = 0

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            # Unchanged paragraphs
            for k in range(i1, i2):
                para = orig_paras[k]
                doc_para = doc.add_paragraph(para["text"])
                # Try to preserve style
                try:
                    doc_para.style = doc.styles[para["style"]]
                except KeyError:
                    pass

        elif tag == "replace":
            # Modified paragraphs
            max_pairs = max(i2 - i1, j2 - j1)
            for k in range(max_pairs):
                old_para = orig_paras[i1 + k] if i1 + k < i2 else None
                new_para = rev_paras[j1 + k] if j1 + k < j2 else None

                old_text = old_para["text"] if old_para else ""
                new_text = new_para["text"] if new_para else ""
                style_name = (old_para or new_para)["style"]

                change_id += 1

                if old_text and new_text:
                    # Build tracked-changes paragraph
                    p_elem = build_tracked_paragraph(old_text, new_text, style_name)
                    doc.element.body.append(p_elem)

                    classification = classify_change(old_text, new_text)
                    section = (old_para or new_para).get("section", "Unknown")

                    changes.append({
                        "id": change_id,
                        "type": "modification",
                        "section": section,
                        "old_text": old_text,
                        "new_text": new_text,
                        **classification,
                    })
                elif old_text:
                    # Deleted paragraph
                    p_elem = build_tracked_paragraph(old_text, "", style_name)
                    doc.element.body.append(p_elem)

                    classification = classify_change(old_text, "")
                    changes.append({
                        "id": change_id,
                        "type": "deletion",
                        "section": old_para.get("section", "Unknown"),
                        "old_text": old_text,
                        "new_text": "",
                        **classification,
                    })
                else:
                    # Added paragraph
                    p_elem = build_tracked_paragraph("", new_text, style_name)
                    doc.element.body.append(p_elem)

                    classification = classify_change("", new_text)
                    changes.append({
                        "id": change_id,
                        "type": "addition",
                        "section": new_para.get("section", "Unknown"),
                        "old_text": "",
                        "new_text": new_text,
                        **classification,
                    })

        elif tag == "delete":
            # Deleted paragraphs
            for k in range(i1, i2):
                para = orig_paras[k]
                change_id += 1

                p_elem = build_tracked_paragraph(para["text"], "", para["style"])
                doc.element.body.append(p_elem)

                classification = classify_change(para["text"], "")
                changes.append({
                    "id": change_id,
                    "type": "deletion",
                    "section": para.get("section", "Unknown"),
                    "old_text": para["text"],
                    "new_text": "",
                    **classification,
                })

        elif tag == "insert":
            # Added paragraphs
            for k in range(j1, j2):
                para = rev_paras[k]
                change_id += 1

                p_elem = build_tracked_paragraph("", para["text"], para["style"])
                doc.element.body.append(p_elem)

                classification = classify_change("", para["text"])
                changes.append({
                    "id": change_id,
                    "type": "addition",
                    "section": para.get("section", "Unknown"),
                    "old_text": "",
                    "new_text": para["text"],
                    **classification,
                })

    doc.save(output_path)
    return changes


# ---------------------------------------------------------------------------
# Output generation
# ---------------------------------------------------------------------------
def write_outputs(changes: list[dict], output_dir: str,
                  original_path: str, revised_path: str) -> dict:
    """Write analysis output files and return summary."""
    os.makedirs(output_dir, exist_ok=True)

    # Stats
    total = len(changes)
    additions = sum(1 for c in changes if c["type"] == "addition")
    deletions = sum(1 for c in changes if c["type"] == "deletion")
    modifications = sum(1 for c in changes if c["type"] == "modification")

    high_risk = [c for c in changes if c["risk"] == "HIGH"]
    medium_risk = [c for c in changes if c["risk"] == "MEDIUM"]
    low_risk = [c for c in changes if c["risk"] == "LOW"]

    substantive = sum(1 for c in changes if c["category"] == "Substantive")
    administrative = sum(1 for c in changes if c["category"] == "Administrative")
    risk_relevant = sum(1 for c in changes if c["category"] == "Risk-relevant")

    # 1. change_analysis.json
    analysis = {
        "original": os.path.abspath(original_path),
        "revised": os.path.abspath(revised_path),
        "generated": REV_DATE,
        "statistics": {
            "total_changes": total,
            "by_type": {
                "additions": additions,
                "deletions": deletions,
                "modifications": modifications,
            },
            "by_risk": {
                "HIGH": len(high_risk),
                "MEDIUM": len(medium_risk),
                "LOW": len(low_risk),
            },
            "by_category": {
                "Substantive": substantive,
                "Administrative": administrative,
                "Risk-relevant": risk_relevant,
            },
        },
        "changes": [
            {
                "id": c["id"],
                "type": c["type"],
                "section": c["section"],
                "category": c["category"],
                "risk": c["risk"],
                "risk_reasons": c["risk_reasons"],
                "old_text": c["old_text"][:500],
                "new_text": c["new_text"][:500],
            }
            for c in changes
        ],
    }

    json_path = os.path.join(output_dir, "change_analysis.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(analysis, f, indent=2)

    # 2. material_changes.txt (HIGH risk only)
    material_lines = [
        "Material Changes (HIGH Risk)",
        "=" * 50,
        "",
        f"Total HIGH-risk changes: {len(high_risk)}",
        "",
    ]

    for c in high_risk:
        material_lines.append(f"Change #{c['id']} [{c['type'].upper()}]")
        material_lines.append(f"  Section: {c['section'][:60]}")
        material_lines.append(f"  Category: {c['category']}")
        material_lines.append(f"  Risk reasons: {'; '.join(c['risk_reasons'])}")
        if c["old_text"]:
            material_lines.append(f"  Old: {c['old_text'][:200]}")
        if c["new_text"]:
            material_lines.append(f"  New: {c['new_text'][:200]}")
        material_lines.append("")

    if not high_risk:
        material_lines.append("No HIGH-risk changes detected.")

    material_path = os.path.join(output_dir, "material_changes.txt")
    with open(material_path, "w", encoding="utf-8") as f:
        f.write("\n".join(material_lines))

    # 3. redline_summary.txt
    summary_lines = [
        "Redline Summary",
        "=" * 50,
        "",
        f"Original: {Path(original_path).name}",
        f"Revised:  {Path(revised_path).name}",
        f"Generated: {REV_DATE}",
        "",
        "Change Statistics",
        "-" * 30,
        f"Total changes:     {total}",
        f"  Additions:       {additions}",
        f"  Deletions:       {deletions}",
        f"  Modifications:   {modifications}",
        "",
        "By Risk Level",
        "-" * 30,
        f"  HIGH:   {len(high_risk)}",
        f"  MEDIUM: {len(medium_risk)}",
        f"  LOW:    {len(low_risk)}",
        "",
        "By Category",
        "-" * 30,
        f"  Substantive:     {substantive}",
        f"  Administrative:  {administrative}",
        f"  Risk-relevant:   {risk_relevant}",
        "",
    ]

    # List sections with changes
    section_counts = {}
    for c in changes:
        s = c["section"]
        section_counts[s] = section_counts.get(s, 0) + 1

    if section_counts:
        summary_lines.append("Changes by Section")
        summary_lines.append("-" * 30)
        for section, count in sorted(section_counts.items(), key=lambda x: x[1], reverse=True):
            summary_lines.append(f"  {section[:60]}: {count}")

    summary_path = os.path.join(output_dir, "redline_summary.txt")
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write("\n".join(summary_lines))

    # Return summary for JSON stdout
    return {
        "status": "success",
        "output_dir": output_dir,
        "original": os.path.abspath(original_path),
        "revised": os.path.abspath(revised_path),
        "total_changes": total,
        "by_type": {"additions": additions, "deletions": deletions, "modifications": modifications},
        "by_risk": {"HIGH": len(high_risk), "MEDIUM": len(medium_risk), "LOW": len(low_risk)},
        "by_category": {"Substantive": substantive, "Administrative": administrative, "Risk-relevant": risk_relevant},
        "high_risk_changes": [
            {
                "id": c["id"],
                "type": c["type"],
                "section": c["section"][:60],
                "risk_reasons": c["risk_reasons"],
                "old_text": c["old_text"][:200],
                "new_text": c["new_text"][:200],
            }
            for c in high_risk[:10]  # Top 10 for stdout
        ],
        "outputs": {
            "redline_docx": os.path.join(output_dir, "redline.docx"),
            "change_analysis_json": json_path,
            "material_changes_txt": material_path,
            "redline_summary_txt": summary_path,
        },
    }


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Generate tracked-changes redline from two contract versions (.docx)."
    )
    parser.add_argument("--original", required=True,
                        help="Path to the original contract (.docx)")
    parser.add_argument("--revised", required=True,
                        help="Path to the revised contract (.docx)")
    parser.add_argument("--output-dir", required=True,
                        help="Directory for redline output")
    args = parser.parse_args()

    original_path = os.path.abspath(args.original)
    revised_path = os.path.abspath(args.revised)
    output_dir = os.path.abspath(args.output_dir)

    # Validate inputs
    if not os.path.isfile(original_path):
        print(json.dumps({"error": f"File not found: {original_path}"}))
        sys.exit(1)
    if not os.path.isfile(revised_path):
        print(json.dumps({"error": f"File not found: {revised_path}"}))
        sys.exit(1)

    if not original_path.lower().endswith(".docx"):
        print(json.dumps({"error": f"Original file must be .docx, got: {Path(original_path).suffix}"}))
        sys.exit(1)
    if not revised_path.lower().endswith(".docx"):
        print(json.dumps({"error": f"Revised file must be .docx, got: {Path(revised_path).suffix}"}))
        sys.exit(1)

    # Load documents
    print(f"Loading original: {Path(original_path).name}", file=sys.stderr)
    try:
        orig_doc = DocxDocument(original_path)
    except Exception as e:
        print(json.dumps({"error": f"Cannot open original document: {e}"}))
        sys.exit(1)

    print(f"Loading revised: {Path(revised_path).name}", file=sys.stderr)
    try:
        rev_doc = DocxDocument(revised_path)
    except Exception as e:
        print(json.dumps({"error": f"Cannot open revised document: {e}"}))
        sys.exit(1)

    # Extract paragraphs
    print("Extracting paragraphs...", file=sys.stderr)
    orig_paras = extract_paragraphs(orig_doc)
    rev_paras = extract_paragraphs(rev_doc)

    print(f"  Original: {len(orig_paras)} paragraphs", file=sys.stderr)
    print(f"  Revised:  {len(rev_paras)} paragraphs", file=sys.stderr)

    # Filter out empty paragraphs for comparison but keep for structure
    orig_nonempty = [p for p in orig_paras if p["text"]]
    rev_nonempty = [p for p in rev_paras if p["text"]]

    if not orig_nonempty:
        print(json.dumps({"error": "No text found in original document"}))
        sys.exit(1)
    if not rev_nonempty:
        print(json.dumps({"error": "No text found in revised document"}))
        sys.exit(1)

    # Check for identical documents
    if [p["text"] for p in orig_nonempty] == [p["text"] for p in rev_nonempty]:
        result = {
            "status": "success",
            "output_dir": output_dir,
            "total_changes": 0,
            "message": "Documents are identical. No differences found.",
        }
        print(json.dumps(result, indent=2))
        sys.exit(0)

    # Assign sections
    orig_nonempty = assign_sections(orig_nonempty)
    rev_nonempty = assign_sections(rev_nonempty)

    # Generate redline document
    os.makedirs(output_dir, exist_ok=True)
    redline_path = os.path.join(output_dir, "redline.docx")

    print("Generating redline with tracked changes...", file=sys.stderr)
    changes = generate_redline_docx(orig_nonempty, rev_nonempty, redline_path)
    print(f"  {len(changes)} changes marked", file=sys.stderr)

    # Write analysis outputs
    print("Writing analysis outputs...", file=sys.stderr)
    summary = write_outputs(changes, output_dir, original_path, revised_path)

    # Print JSON to stdout for Claude to parse
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
