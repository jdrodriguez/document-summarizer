#!/usr/bin/env python3
"""
Document comparison script for the legal-doc-compare skill.

Takes two documents (DOCX, PDF, or TXT) and generates a comprehensive
comparison with visual diffs, change heatmaps, and structured change logs.

Usage:
    python3 compare_documents.py --file1 <path> --file2 <path> \
        --output-dir <dir> [--labels "Original,Revised"]
"""
import argparse
import difflib
import html
import json
import os
import re
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Dependency imports (with graceful fallbacks)
# ---------------------------------------------------------------------------
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import plotly.graph_objects as go
    import plotly.io as pio
except ImportError:
    go = None
    pio = None


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------
def extract_text_docx(filepath: str) -> list[str]:
    """Extract paragraphs from a DOCX file."""
    if DocxDocument is None:
        print("python-docx not available", file=sys.stderr)
        return []
    try:
        doc = DocxDocument(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(c for c in cells if c)
                if row_text:
                    paragraphs.append(row_text)
        return paragraphs
    except Exception as e:
        print(f"DOCX extraction failed: {e}", file=sys.stderr)
        return []


def extract_text_pdf(filepath: str) -> list[str]:
    """Extract paragraphs from a PDF file."""
    if pdfplumber is None:
        print("pdfplumber not available", file=sys.stderr)
        return []
    try:
        paragraphs = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    # Split into paragraphs on double newlines
                    for para in re.split(r'\n\s*\n', text):
                        para = para.strip()
                        if para:
                            # Collapse internal whitespace
                            para = re.sub(r'\s+', ' ', para)
                            paragraphs.append(para)
        return paragraphs
    except Exception as e:
        print(f"PDF extraction failed: {e}", file=sys.stderr)
        return []


def extract_text_txt(filepath: str) -> list[str]:
    """Extract paragraphs from a text file."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        paragraphs = []
        for para in re.split(r'\n\s*\n', content):
            para = para.strip()
            if para:
                paragraphs.append(para)
        return paragraphs
    except Exception as e:
        print(f"Text extraction failed: {e}", file=sys.stderr)
        return []


def extract_paragraphs(filepath: str) -> list[str]:
    """Route extraction based on file type."""
    ext = Path(filepath).suffix.lower()
    if ext == ".docx":
        return extract_text_docx(filepath)
    elif ext == ".pdf":
        return extract_text_pdf(filepath)
    elif ext == ".txt":
        return extract_text_txt(filepath)
    else:
        print(f"Unsupported file type: {ext}", file=sys.stderr)
        return []


# ---------------------------------------------------------------------------
# Section detection
# ---------------------------------------------------------------------------
SECTION_RE = re.compile(
    r'^(?:(?:ARTICLE|SECTION|CHAPTER|PART)\s+[\dIVXLCDM]+|'
    r'(?:\d+\.)+\s+|'
    r'[IVXLCDM]+\.\s+|'
    r'[A-Z]\.\s+)',
    re.IGNORECASE
)


def detect_section(text: str) -> str:
    """Try to extract a section heading from paragraph text."""
    if SECTION_RE.match(text):
        # Return first 80 chars as section name
        return text[:80].strip()
    # Check for all-caps short lines (likely headings)
    if len(text) < 100 and text == text.upper() and len(text.split()) >= 2:
        return text[:80].strip()
    return ""


def assign_sections(paragraphs: list[str]) -> list[dict]:
    """Assign each paragraph to a section."""
    result = []
    current_section = "Preamble"
    for i, para in enumerate(paragraphs):
        section_name = detect_section(para)
        if section_name:
            current_section = section_name
        result.append({
            "index": i,
            "text": para,
            "section": current_section,
        })
    return result


# ---------------------------------------------------------------------------
# Comparison engine
# ---------------------------------------------------------------------------
def compute_paragraph_diff(paras1: list[str], paras2: list[str]) -> list[dict]:
    """
    Compute paragraph-level diff between two document versions.
    Returns a list of change records.
    """
    matcher = difflib.SequenceMatcher(None, paras1, paras2)
    changes = []
    change_id = 0

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue

        if tag == "replace":
            # Modified paragraphs -- do word-level diff within each pair
            max_pairs = max(i2 - i1, j2 - j1)
            for k in range(max_pairs):
                old_idx = i1 + k if i1 + k < i2 else None
                new_idx = j1 + k if j1 + k < j2 else None
                old_text = paras1[old_idx] if old_idx is not None else ""
                new_text = paras2[new_idx] if new_idx is not None else ""

                change_id += 1

                if old_text and new_text:
                    # Word-level diff
                    word_changes = compute_word_diff(old_text, new_text)
                    changes.append({
                        "id": change_id,
                        "type": "modification",
                        "old_text": old_text,
                        "new_text": new_text,
                        "old_index": old_idx,
                        "new_index": new_idx,
                        "word_changes": word_changes,
                    })
                elif old_text:
                    changes.append({
                        "id": change_id,
                        "type": "deletion",
                        "old_text": old_text,
                        "new_text": "",
                        "old_index": old_idx,
                        "new_index": None,
                        "word_changes": [],
                    })
                else:
                    changes.append({
                        "id": change_id,
                        "type": "addition",
                        "old_text": "",
                        "new_text": new_text,
                        "old_index": None,
                        "new_index": new_idx,
                        "word_changes": [],
                    })

        elif tag == "delete":
            for k in range(i1, i2):
                change_id += 1
                changes.append({
                    "id": change_id,
                    "type": "deletion",
                    "old_text": paras1[k],
                    "new_text": "",
                    "old_index": k,
                    "new_index": None,
                    "word_changes": [],
                })

        elif tag == "insert":
            for k in range(j1, j2):
                change_id += 1
                changes.append({
                    "id": change_id,
                    "type": "addition",
                    "old_text": "",
                    "new_text": paras2[k],
                    "old_index": None,
                    "new_index": k,
                    "word_changes": [],
                })

    return changes


def compute_word_diff(old_text: str, new_text: str) -> list[dict]:
    """Compute word-level diff between two strings."""
    old_words = old_text.split()
    new_words = new_text.split()
    matcher = difflib.SequenceMatcher(None, old_words, new_words)
    changes = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        changes.append({
            "type": tag,
            "old_words": " ".join(old_words[i1:i2]),
            "new_words": " ".join(new_words[j1:j2]),
        })

    return changes


# ---------------------------------------------------------------------------
# Section-level change mapping
# ---------------------------------------------------------------------------
def map_changes_to_sections(changes: list[dict], sections1: list[dict],
                            sections2: list[dict]) -> dict:
    """Map changes to their document sections and compute per-section stats."""
    section_changes = {}

    for change in changes:
        section = "Unknown"
        if change["old_index"] is not None and change["old_index"] < len(sections1):
            section = sections1[change["old_index"]]["section"]
        elif change["new_index"] is not None and change["new_index"] < len(sections2):
            section = sections2[change["new_index"]]["section"]

        change["section"] = section

        if section not in section_changes:
            section_changes[section] = {
                "additions": 0,
                "deletions": 0,
                "modifications": 0,
                "total": 0,
            }

        section_changes[section][change["type"] + "s" if not change["type"].endswith("s") else change["type"]] = \
            section_changes[section].get(change["type"] + "s" if not change["type"].endswith("s") else change["type"], 0) + 1
        section_changes[section]["total"] += 1

    # Fix the key names (handle pluralization properly)
    for section in section_changes:
        stats = section_changes[section]
        # Normalize keys
        normalized = {"additions": 0, "deletions": 0, "modifications": 0, "total": stats["total"]}
        for key, val in stats.items():
            if "addition" in key:
                normalized["additions"] += val
            elif "deletion" in key:
                normalized["deletions"] += val
            elif "modification" in key:
                normalized["modifications"] += val
        section_changes[section] = normalized

    return section_changes


# ---------------------------------------------------------------------------
# HTML comparison output
# ---------------------------------------------------------------------------
def generate_comparison_html(paras1: list[str], paras2: list[str],
                             changes: list[dict], label1: str, label2: str) -> str:
    """Generate a self-contained side-by-side HTML comparison."""
    # Build maps for quick lookup
    old_changes = {}  # old_index -> change
    new_changes = {}  # new_index -> change
    for change in changes:
        if change["old_index"] is not None:
            old_changes[change["old_index"]] = change
        if change["new_index"] is not None:
            new_changes[change["new_index"]] = change

    # Build aligned rows using SequenceMatcher
    matcher = difflib.SequenceMatcher(None, paras1, paras2)
    rows = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for k in range(i2 - i1):
                rows.append(("equal", paras1[i1 + k], paras2[j1 + k]))
        elif tag == "replace":
            max_len = max(i2 - i1, j2 - j1)
            for k in range(max_len):
                old = paras1[i1 + k] if i1 + k < i2 else None
                new = paras2[j1 + k] if j1 + k < j2 else None
                if old and new:
                    rows.append(("modified", old, new))
                elif old:
                    rows.append(("deleted", old, None))
                else:
                    rows.append(("added", None, new))
        elif tag == "delete":
            for k in range(i1, i2):
                rows.append(("deleted", paras1[k], None))
        elif tag == "insert":
            for k in range(j1, j2):
                rows.append(("added", None, paras2[k]))

    # Generate HTML rows
    html_rows = []
    for row_type, left, right in rows:
        left_html = html.escape(left) if left else ""
        right_html = html.escape(right) if right else ""

        if row_type == "equal":
            html_rows.append(f'''
                <tr>
                    <td class="equal">{left_html}</td>
                    <td class="equal">{right_html}</td>
                </tr>''')
        elif row_type == "modified":
            # Word-level highlighting
            left_highlighted = highlight_word_diff(left, right, "old")
            right_highlighted = highlight_word_diff(left, right, "new")
            html_rows.append(f'''
                <tr>
                    <td class="modified">{left_highlighted}</td>
                    <td class="modified">{right_highlighted}</td>
                </tr>''')
        elif row_type == "deleted":
            html_rows.append(f'''
                <tr>
                    <td class="deleted">{left_html}</td>
                    <td class="empty"></td>
                </tr>''')
        elif row_type == "added":
            html_rows.append(f'''
                <tr>
                    <td class="empty"></td>
                    <td class="added">{right_html}</td>
                </tr>''')

    rows_html = "\n".join(html_rows)

    # Stats
    additions = sum(1 for c in changes if c["type"] == "addition")
    deletions = sum(1 for c in changes if c["type"] == "deletion")
    modifications = sum(1 for c in changes if c["type"] == "modification")

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document Comparison: {html.escape(label1)} vs {html.escape(label2)}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; background: #f5f5f5; padding: 20px; }}
        .header {{ background: #1a1a2e; color: white; padding: 20px 30px; border-radius: 8px 8px 0 0; }}
        .header h1 {{ font-size: 1.4em; margin-bottom: 10px; }}
        .stats {{ display: flex; gap: 20px; font-size: 0.9em; }}
        .stat {{ padding: 4px 12px; border-radius: 4px; }}
        .stat-add {{ background: #22c55e33; color: #16a34a; }}
        .stat-del {{ background: #ef444433; color: #dc2626; }}
        .stat-mod {{ background: #eab30833; color: #ca8a04; }}
        .container {{ background: white; border-radius: 0 0 8px 8px; overflow: hidden; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }}
        table {{ width: 100%; border-collapse: collapse; table-layout: fixed; }}
        thead th {{ background: #e5e7eb; padding: 12px 16px; text-align: left; font-weight: 600; font-size: 0.85em; text-transform: uppercase; letter-spacing: 0.05em; color: #374151; position: sticky; top: 0; }}
        td {{ padding: 10px 16px; vertical-align: top; font-size: 0.9em; line-height: 1.6; border-bottom: 1px solid #f3f4f6; word-wrap: break-word; }}
        .equal {{ color: #374151; }}
        .deleted {{ background: #fef2f2; color: #991b1b; }}
        .added {{ background: #f0fdf4; color: #166534; }}
        .modified {{ background: #fffbeb; }}
        .empty {{ background: #f9fafb; }}
        .word-del {{ background: #fca5a5; text-decoration: line-through; padding: 1px 2px; border-radius: 2px; }}
        .word-add {{ background: #86efac; padding: 1px 2px; border-radius: 2px; }}
        .legend {{ padding: 16px 20px; background: #f9fafb; border-top: 1px solid #e5e7eb; font-size: 0.85em; color: #6b7280; }}
        .legend span {{ margin-right: 20px; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Document Comparison</h1>
        <div class="stats">
            <span class="stat stat-add">{additions} additions</span>
            <span class="stat stat-del">{deletions} deletions</span>
            <span class="stat stat-mod">{modifications} modifications</span>
            <span class="stat">{len(changes)} total changes</span>
        </div>
    </div>
    <div class="container">
        <table>
            <thead>
                <tr>
                    <th style="width:50%">{html.escape(label1)}</th>
                    <th style="width:50%">{html.escape(label2)}</th>
                </tr>
            </thead>
            <tbody>
                {rows_html}
            </tbody>
        </table>
        <div class="legend">
            <span style="color:#16a34a">Green = Added</span>
            <span style="color:#dc2626">Red = Deleted</span>
            <span style="color:#ca8a04">Yellow = Modified</span>
        </div>
    </div>
</body>
</html>"""


def highlight_word_diff(old_text: str, new_text: str, side: str) -> str:
    """Generate HTML with word-level highlighting."""
    old_words = old_text.split()
    new_words = new_text.split()
    matcher = difflib.SequenceMatcher(None, old_words, new_words)
    result_parts = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            if side == "old":
                result_parts.append(html.escape(" ".join(old_words[i1:i2])))
            else:
                result_parts.append(html.escape(" ".join(new_words[j1:j2])))
        elif tag == "replace":
            if side == "old":
                result_parts.append(
                    f'<span class="word-del">{html.escape(" ".join(old_words[i1:i2]))}</span>'
                )
            else:
                result_parts.append(
                    f'<span class="word-add">{html.escape(" ".join(new_words[j1:j2]))}</span>'
                )
        elif tag == "delete":
            if side == "old":
                result_parts.append(
                    f'<span class="word-del">{html.escape(" ".join(old_words[i1:i2]))}</span>'
                )
        elif tag == "insert":
            if side == "new":
                result_parts.append(
                    f'<span class="word-add">{html.escape(" ".join(new_words[j1:j2]))}</span>'
                )

    return " ".join(result_parts)


# ---------------------------------------------------------------------------
# Heatmap generation
# ---------------------------------------------------------------------------
def generate_heatmap_html(section_changes: dict, label1: str, label2: str) -> str:
    """Generate a plotly change density heatmap as self-contained HTML."""
    if go is None or pio is None:
        return generate_heatmap_fallback(section_changes, label1, label2)

    sections = list(section_changes.keys())
    additions = [section_changes[s]["additions"] for s in sections]
    deletions = [section_changes[s]["deletions"] for s in sections]
    modifications = [section_changes[s]["modifications"] for s in sections]
    totals = [section_changes[s]["total"] for s in sections]

    # Truncate long section names for display
    display_names = [s[:60] + "..." if len(s) > 60 else s for s in sections]

    fig = go.Figure()

    fig.add_trace(go.Bar(
        name="Additions",
        x=display_names,
        y=additions,
        marker_color="#22c55e",
    ))
    fig.add_trace(go.Bar(
        name="Deletions",
        x=display_names,
        y=deletions,
        marker_color="#ef4444",
    ))
    fig.add_trace(go.Bar(
        name="Modifications",
        x=display_names,
        y=modifications,
        marker_color="#eab308",
    ))

    fig.update_layout(
        title=f"Change Density by Section: {label1} vs {label2}",
        barmode="stack",
        xaxis_title="Section",
        yaxis_title="Number of Changes",
        template="plotly_white",
        height=max(400, len(sections) * 30 + 200),
        font=dict(family="-apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif"),
    )

    if len(sections) > 5:
        fig.update_layout(xaxis_tickangle=-45)

    return pio.to_html(fig, full_html=True, include_plotlyjs=True)


def generate_heatmap_fallback(section_changes: dict, label1: str, label2: str) -> str:
    """Generate a simple HTML bar chart without plotly."""
    sections = sorted(section_changes.keys(), key=lambda s: section_changes[s]["total"], reverse=True)
    max_total = max((section_changes[s]["total"] for s in sections), default=1)

    bars_html = ""
    for section in sections:
        stats = section_changes[section]
        pct = (stats["total"] / max_total) * 100 if max_total > 0 else 0
        add_pct = (stats["additions"] / max_total) * 100 if max_total > 0 else 0
        del_pct = (stats["deletions"] / max_total) * 100 if max_total > 0 else 0
        mod_pct = (stats["modifications"] / max_total) * 100 if max_total > 0 else 0

        display = html.escape(section[:60])
        bars_html += f'''
        <div class="bar-row">
            <div class="bar-label" title="{html.escape(section)}">{display}</div>
            <div class="bar-track">
                <div class="bar-fill bar-add" style="width:{add_pct}%"></div>
                <div class="bar-fill bar-del" style="width:{del_pct}%"></div>
                <div class="bar-fill bar-mod" style="width:{mod_pct}%"></div>
            </div>
            <div class="bar-count">{stats["total"]}</div>
        </div>'''

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Change Heatmap: {html.escape(label1)} vs {html.escape(label2)}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, sans-serif; padding: 30px; background: #f5f5f5; }}
        h1 {{ font-size: 1.3em; margin-bottom: 20px; color: #1a1a2e; }}
        .bar-row {{ display: flex; align-items: center; margin-bottom: 8px; }}
        .bar-label {{ width: 250px; font-size: 0.85em; color: #374151; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }}
        .bar-track {{ flex: 1; height: 24px; background: #e5e7eb; border-radius: 4px; display: flex; overflow: hidden; }}
        .bar-fill {{ height: 100%; }}
        .bar-add {{ background: #22c55e; }}
        .bar-del {{ background: #ef4444; }}
        .bar-mod {{ background: #eab308; }}
        .bar-count {{ width: 40px; text-align: right; font-size: 0.85em; color: #6b7280; margin-left: 8px; }}
        .legend {{ margin-top: 20px; font-size: 0.85em; color: #6b7280; }}
        .legend span {{ margin-right: 20px; }}
    </style>
</head>
<body>
    <h1>Change Density by Section</h1>
    {bars_html}
    <div class="legend">
        <span style="color:#22c55e">Green = Additions</span>
        <span style="color:#ef4444">Red = Deletions</span>
        <span style="color:#eab308">Yellow = Modifications</span>
    </div>
</body>
</html>"""


# ---------------------------------------------------------------------------
# Output generation
# ---------------------------------------------------------------------------
def write_outputs(paras1: list[str], paras2: list[str], changes: list[dict],
                  section_changes: dict, output_dir: str, label1: str,
                  label2: str, file1: str, file2: str) -> dict:
    """Write all output files and return summary."""
    os.makedirs(output_dir, exist_ok=True)

    # Stats
    additions = sum(1 for c in changes if c["type"] == "addition")
    deletions = sum(1 for c in changes if c["type"] == "deletion")
    modifications = sum(1 for c in changes if c["type"] == "modification")
    total_paras = max(len(paras1), len(paras2))
    pct_changed = round(len(changes) / total_paras * 100, 1) if total_paras > 0 else 0

    # Most changed sections
    most_changed = sorted(
        section_changes.items(),
        key=lambda x: x[1]["total"],
        reverse=True
    )[:5]

    # 1. comparison.html
    print("  Generating comparison.html...", file=sys.stderr)
    comparison_html = generate_comparison_html(paras1, paras2, changes, label1, label2)
    html_path = os.path.join(output_dir, "comparison.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(comparison_html)

    # 2. change_log.json
    change_log = {
        "file1": os.path.abspath(file1),
        "file2": os.path.abspath(file2),
        "label1": label1,
        "label2": label2,
        "statistics": {
            "total_changes": len(changes),
            "additions": additions,
            "deletions": deletions,
            "modifications": modifications,
            "paragraphs_in_original": len(paras1),
            "paragraphs_in_revised": len(paras2),
            "percentage_changed": pct_changed,
        },
        "section_changes": section_changes,
        "changes": [
            {
                "id": c["id"],
                "type": c["type"],
                "section": c.get("section", "Unknown"),
                "old_text": c["old_text"],
                "new_text": c["new_text"],
                "word_changes": c.get("word_changes", []),
            }
            for c in changes
        ],
    }

    json_path = os.path.join(output_dir, "change_log.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(change_log, f, indent=2)

    # 3. change_heatmap.html
    print("  Generating change_heatmap.html...", file=sys.stderr)
    if section_changes:
        heatmap_html = generate_heatmap_html(section_changes, label1, label2)
        heatmap_path = os.path.join(output_dir, "change_heatmap.html")
        with open(heatmap_path, "w", encoding="utf-8") as f:
            f.write(heatmap_html)

    # 4. comparison_summary.txt
    summary_lines = [
        "Document Comparison Summary",
        "=" * 50,
        "",
        f"Original: {Path(file1).name} ({label1})",
        f"Revised:  {Path(file2).name} ({label2})",
        "",
        "Change Statistics",
        "-" * 30,
        f"Total changes:     {len(changes)}",
        f"  Additions:       {additions}",
        f"  Deletions:       {deletions}",
        f"  Modifications:   {modifications}",
        f"",
        f"Original paragraphs: {len(paras1)}",
        f"Revised paragraphs:  {len(paras2)}",
        f"Percentage changed:  {pct_changed}%",
        "",
        "Most Changed Sections",
        "-" * 30,
    ]

    for section_name, stats in most_changed:
        summary_lines.append(
            f"  {section_name[:60]}: {stats['total']} changes "
            f"({stats['additions']} add, {stats['deletions']} del, {stats['modifications']} mod)"
        )

    summary_lines.extend(["", "Top 5 Largest Changes", "-" * 30])
    # Sort changes by text length difference
    sorted_changes = sorted(
        changes,
        key=lambda c: max(len(c["old_text"]), len(c["new_text"])),
        reverse=True
    )[:5]

    for c in sorted_changes:
        summary_lines.append(f"  [{c['type'].upper()}] Section: {c.get('section', 'Unknown')[:40]}")
        if c["old_text"]:
            summary_lines.append(f"    Old: {c['old_text'][:120]}...")
        if c["new_text"]:
            summary_lines.append(f"    New: {c['new_text'][:120]}...")
        summary_lines.append("")

    summary_path = os.path.join(output_dir, "comparison_summary.txt")
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write("\n".join(summary_lines))

    # Return summary for JSON stdout output
    return {
        "status": "success",
        "output_dir": output_dir,
        "file1": os.path.abspath(file1),
        "file2": os.path.abspath(file2),
        "label1": label1,
        "label2": label2,
        "total_changes": len(changes),
        "additions": additions,
        "deletions": deletions,
        "modifications": modifications,
        "percentage_changed": pct_changed,
        "paragraphs_original": len(paras1),
        "paragraphs_revised": len(paras2),
        "most_changed_sections": [
            {"section": name[:80], "total": stats["total"]}
            for name, stats in most_changed
        ],
        "outputs": {
            "comparison_html": html_path,
            "change_log_json": json_path,
            "change_heatmap_html": os.path.join(output_dir, "change_heatmap.html"),
            "comparison_summary": summary_path,
        },
    }


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Compare two documents (DOCX, PDF, TXT) and generate "
                    "visual diffs, change heatmaps, and structured change logs."
    )
    parser.add_argument("--file1", required=True,
                        help="Path to the original document")
    parser.add_argument("--file2", required=True,
                        help="Path to the revised document")
    parser.add_argument("--output-dir", required=True,
                        help="Directory for comparison output")
    parser.add_argument("--labels", default="Original,Revised",
                        help="Comma-separated labels for file1 and file2 (default: Original,Revised)")
    args = parser.parse_args()

    file1 = os.path.abspath(args.file1)
    file2 = os.path.abspath(args.file2)
    output_dir = os.path.abspath(args.output_dir)

    labels = args.labels.split(",", 1)
    label1 = labels[0].strip() if len(labels) > 0 else "Original"
    label2 = labels[1].strip() if len(labels) > 1 else "Revised"

    # Validate inputs
    if not os.path.isfile(file1):
        print(json.dumps({"error": f"File not found: {file1}"}))
        sys.exit(1)
    if not os.path.isfile(file2):
        print(json.dumps({"error": f"File not found: {file2}"}))
        sys.exit(1)

    supported = {".pdf", ".docx", ".txt"}
    ext1 = Path(file1).suffix.lower()
    ext2 = Path(file2).suffix.lower()
    if ext1 not in supported:
        print(json.dumps({"error": f"Unsupported file type: {ext1}", "supported": list(supported)}))
        sys.exit(1)
    if ext2 not in supported:
        print(json.dumps({"error": f"Unsupported file type: {ext2}", "supported": list(supported)}))
        sys.exit(1)

    # Extract text
    print(f"Extracting text from {Path(file1).name}...", file=sys.stderr)
    paras1 = extract_paragraphs(file1)
    if not paras1:
        print(json.dumps({"error": f"No text extracted from {Path(file1).name}. File may be scanned -- try /legal-ocr:process first."}))
        sys.exit(1)

    print(f"Extracting text from {Path(file2).name}...", file=sys.stderr)
    paras2 = extract_paragraphs(file2)
    if not paras2:
        print(json.dumps({"error": f"No text extracted from {Path(file2).name}. File may be scanned -- try /legal-ocr:process first."}))
        sys.exit(1)

    print(f"  {Path(file1).name}: {len(paras1)} paragraphs", file=sys.stderr)
    print(f"  {Path(file2).name}: {len(paras2)} paragraphs", file=sys.stderr)

    # Check for identical documents
    if paras1 == paras2:
        result = {
            "status": "success",
            "output_dir": output_dir,
            "total_changes": 0,
            "message": "Documents are identical. No differences found.",
        }
        print(json.dumps(result, indent=2))
        sys.exit(0)

    # Compute diffs
    print("Computing paragraph-level diff...", file=sys.stderr)
    changes = compute_paragraph_diff(paras1, paras2)
    print(f"  {len(changes)} changes found", file=sys.stderr)

    # Map to sections
    sections1 = assign_sections(paras1)
    sections2 = assign_sections(paras2)
    section_changes = map_changes_to_sections(changes, sections1, sections2)

    # Write outputs
    print("Generating outputs...", file=sys.stderr)
    summary = write_outputs(paras1, paras2, changes, section_changes,
                            output_dir, label1, label2, file1, file2)

    # Print JSON to stdout for Claude to parse
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
